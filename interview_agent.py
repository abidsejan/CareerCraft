"""
╔══════════════════════════════════════════════════════════════════╗
║   CareerCraft Pro  —  World-Class Career Acceleration Suite      ║
║   Interview Coach · CV Analyser · Resume Rewriter · Cover Letter ║
╚══════════════════════════════════════════════════════════════════╝

pip install streamlit groq PyPDF2 python-docx requests beautifulsoup4

streamlit run interview_agent.py
"""

import streamlit as st
import json, os, re, io, csv, textwrap
from datetime import datetime
from groq import Groq

try:
    import PyPDF2
    PDF_SUPPORT = True
except ImportError:
    PDF_SUPPORT = False

try:
    from docx import Document as DocxDocument
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX = True
except ImportError:
    DOCX = False

try:
    import requests
    from bs4 import BeautifulSoup
    SCRAPE = True
except ImportError:
    SCRAPE = False

# ═══════════════════════════════════════════════════════════════
# CONFIGURATION
# ═══════════════════════════════════════════════════════════════

HISTORY_FILE = "careercraft_history.json"
GROQ_MODEL   = "llama-3.3-70b-versatile"
GROQ_API_KEY = st.secrets["GROQ_API_KEY"]

MODES = {
    "Full Interview":  "full",
    "Behavioural":     "behavioural",
    "Technical":       "technical",
    "Situational":     "situational",
    "Speed Round":     "speed",
    "Stress Test":     "stress",
}
DIFFICULTIES = {
    "Junior (0–2 yrs)":    "junior",
    "Mid-level (2–5 yrs)": "mid",
    "Senior (5+ yrs)":     "senior",
    "Executive / Lead":    "executive",
}
INDUSTRIES = [
    "Technology / Software", "Finance / Banking", "Healthcare / Biotech",
    "Product Management", "Data Science / AI / ML", "Design / UX",
    "Marketing / Growth", "Sales / Business Dev", "Operations / Strategy",
    "Legal / Compliance", "Education / Research", "Other",
]
CV_TONES = ["Professional & Confident", "Formal & Conservative", "Modern & Creative"]
LETTER_STYLES = ["Standard Professional", "Bold & Assertive", "Warm & Personalised"]

# ═══════════════════════════════════════════════════════════════
# INTERVIEW SYSTEM PROMPT
# ═══════════════════════════════════════════════════════════════

def build_interview_prompt(mode, difficulty, industry, cv_text, weak_context=""):
    diff = {
        "junior":    "Entry-level candidate. Accept simpler examples. Prioritise learning mindset and potential.",
        "mid":       "Mid-level. Expect 2+ specific examples per answer with quantified outcomes.",
        "senior":    "Senior. Demand strategic thinking, leadership evidence, and measurable business impact.",
        "executive": "Executive. Expect P&L ownership, board-level vision, org transformation narratives.",
    }.get(difficulty, "")
    mode_instr = {
        "full":        "Generate a balanced mix: 2 behavioural, 1 situational, 1 technical/role-specific, 1 motivation question.",
        "behavioural": "Generate behavioural questions only. Every question must start with 'Tell me about a time...' or 'Describe a situation where...' or 'Give me an example of...'",
        "technical":   "Generate technical hard-skill questions only. Test domain expertise and problem-solving depth.",
        "situational": "Generate hypothetical situational questions only. Start with 'What would you do if...' or 'Imagine you are...' or 'How would you handle...'",
        "speed":       "Generate 8 short, punchy questions. Expect crisp 2–3 sentence answers. Fast pace. No padding.",
        "stress":      "You are a notoriously tough interviewer. Push back on every answer. Ask 'Why?' and 'How exactly?' constantly. Challenge assumptions. Make the candidate defend their statements.",
    }.get(mode, "")
    cv_section  = f"\n\nCANDIDATE RESUME:\n{cv_text[:3800]}" if cv_text else "\n\nNo resume provided — generate generic role-appropriate questions."
    weak_section = f"\n\nPERFORMANCE CONTEXT FROM PREVIOUS SESSIONS:\n{weak_context}" if weak_context else ""

    return f"""You are a world-class interview coach and former Head of Talent at a Fortune 500 {industry} company.
You have personally conducted 5,000+ interviews and trained 200+ hiring managers.

INTERVIEW MODE: {mode_instr}
CANDIDATE LEVEL: {diff}
TARGET INDUSTRY: {industry}{cv_section}{weak_section}

═══ WHEN OPENING A SESSION ═══
• Number every question clearly (1. 2. 3. etc.)
• Tag each with its type in brackets: [Behavioural] [Technical] [Situational] [Motivation] [Culture Fit]
• If a resume is provided, make questions hyper-specific — reference actual company names, projects, dates
• After questions, write exactly: "Type your answer to any question. I will grade it immediately."

═══ WHEN GRADING AN ANSWER ═══
Return your feedback in EXACTLY this structure — no exceptions:

━━━ SCORE: X/10 ━━━

WHAT WORKED:
• [Specific strength — reference the candidate's exact words]
• [Second strength if present]

STAR AUDIT:
• Situation ── [Clear / Vague / Missing] — [one line explanation]
• Task ──────── [Clear / Vague / Missing] — [one line explanation]
• Action ─────── [Clear / Vague / Missing] — [one line explanation]  
• Result ─────── [Quantified / Present / Missing] — [one line explanation]

COACHING POINTS:
1. [Concrete, specific improvement — never generic advice like "be more specific"]
2. [Second actionable improvement]

REWRITTEN ANSWER:
"[Craft a polished version of their answer. Use their exact experience from the resume where available. Include specific numbers, percentages, or timelines. Complete all four STAR elements. 5–7 sentences. Make it sound authentically human, not robotic.]"

NEXT QUESTION:
"[One sharp, realistic follow-up a tough interviewer would ask based on their answer]"

═══ TONE ═══
Honest, direct, mentor-like warmth. Celebrate genuine strengths. Never be vague. Always be specific."""


# ═══════════════════════════════════════════════════════════════
# CV ANALYSIS PROMPT
# ═══════════════════════════════════════════════════════════════

CV_ANALYSIS_PROMPT = """You are a Principal Recruiter with 20 years of experience and an ATS specialist.
You have reviewed CVs for Google, McKinsey, Goldman Sachs, and 500+ other top employers.

Analyse the candidate's CV against the job description with surgical precision.
Return EXACTLY this format — every section is mandatory:

MATCH SCORE: XX/100

SECTION SCORES:
- Work Experience: XX/25
- Skills & Keywords: XX/25  
- Education & Credentials: XX/20
- CV Format & ATS: XX/15
- Achievements & Impact: XX/15

KEYWORDS PRESENT: [comma-separated list of JD keywords found in CV]
KEYWORDS MISSING: [comma-separated list of critical JD keywords absent from CV]

STRENGTHS:
• [Specific strength with evidence from the CV]
• [Second strength]
• [Third strength]

CRITICAL GAPS:
• [Gap 1 — be precise about what's missing and why it matters for this role]
• [Gap 2]
• [Gap 3]

QUICK WINS (changes achievable in under 15 minutes):
1. [Specific, actionable change — e.g. "Add 'stakeholder management' to your Skills section — it appears 4 times in the JD"]
2. [Second quick win]
3. [Third quick win]

ATS COMPATIBILITY:
• Format issues: [specific formatting problems that hurt ATS parsing, or "None detected"]
• Keyword density: [assessment of keyword optimisation]
• Section headers: [whether standard headers are used, or custom ones that confuse ATS]

GRAMMAR & LANGUAGE AUDIT:
• Tense consistency: [assessment — CVs must use past tense for past roles, present for current]
• Action verbs: [are they strong? e.g. "Led" vs "Was responsible for"]
• Quantification: [how well results are quantified]
• Issues found: [list specific grammar or language problems, or "No significant issues"]

OVERALL VERDICT:
[3–4 sentence honest assessment. Would this CV pass a 6-second recruiter scan? Would it pass ATS? What is the single most important change this candidate should make?]"""


# ═══════════════════════════════════════════════════════════════
# CV REWRITE PROMPT
# ═══════════════════════════════════════════════════════════════

def build_rewrite_prompt(tone):
    tone_instructions = {
        "Professional & Confident": "Write in a confident, authoritative tone. Use powerful action verbs (Spearheaded, Orchestrated, Pioneered, Delivered). Quantify every achievement. Project confidence without arrogance.",
        "Formal & Conservative": "Write in formal, precise language suitable for finance, law, or government sectors. Conservative vocabulary. Emphasis on accuracy and credentials. Measured and professional.",
        "Modern & Creative": "Write in a fresh, dynamic tone for tech, design, or startup roles. Conversational but professional. Show personality. Emphasise impact and innovation.",
    }.get(tone, "Write professionally and confidently.")

    return f"""You are an award-winning CV writer and career strategist. You have helped 10,000+ candidates land roles at top companies.

TONE INSTRUCTION: {tone_instructions}

REWRITING RULES:
1. NEVER fabricate — keep all real experience, companies, dates, qualifications
2. Rewrite EVERY bullet point to start with a strong action verb
3. Add specific metrics to every achievement bullet — if the original has none, infer reasonable estimates and mark with [approx.]
4. Integrate missing keywords from the JD naturally — never keyword-stuff
5. Rewrite the professional profile/summary to be role-specific and compelling (3–4 punchy sentences)
6. Use consistent past tense for past roles, present tense for current role
7. Remove clichés: "team player", "hard worker", "passionate about", "references available"
8. Grammar-proof every sentence — check subject-verb agreement, tense consistency, punctuation
9. Keep dates in consistent format (Month YYYY or YYYY–YYYY)
10. Make every line earn its place — cut filler, add value

OUTPUT FORMAT — use EXACTLY these section markers in square brackets:
[NAME]
[CONTACT]
[PROFILE]
[EXPERIENCE]
[EDUCATION]
[SKILLS]
[ACHIEVEMENTS]
[CERTIFICATIONS]

Under EXPERIENCE, format every role as:
JOB TITLE — Company Name (Month YYYY – Month YYYY or Present)
• Action verb + what you did + result/impact with number
• Action verb + second responsibility + outcome

Produce a complete, polished, ready-to-submit CV. This must be the best version of this person's career story."""


# ═══════════════════════════════════════════════════════════════
# COVER LETTER PROMPT
# ═══════════════════════════════════════════════════════════════

def build_cover_letter_prompt(style):
    style_instructions = {
        "Standard Professional": "Professional, warm, and balanced. Demonstrates competence without being aggressive. Suitable for most corporate roles.",
        "Bold & Assertive": "Direct, confident, almost provocative. Opens with a bold statement of value. Makes the recruiter sit up. Suitable for sales, leadership, or competitive roles.",
        "Warm & Personalised": "Genuine, human, and personal. References specific things about the company culture and mission. Shows authentic enthusiasm. Suitable for mission-driven organisations.",
    }.get(style, "Write professionally and compellingly.")

    return f"""You are one of the world's best cover letter writers. Your letters consistently generate interview callbacks.

STYLE: {style_instructions}

COVER LETTER RULES:
1. NEVER open with "I am writing to apply for..." or "I believe I would be a great fit" — these are instant delete
2. Opening paragraph: Start with a compelling hook — a specific achievement, a bold statement, or a reference to something specific about the company
3. Body paragraph 1: Your single most relevant achievement, with a number, directly tied to a key JD requirement
4. Body paragraph 2: Cultural fit and second key qualification — show you've researched the company
5. Closing paragraph: Clear, confident call to action. Not begging. Expecting.
6. Total length: Exactly 4 paragraphs, fits on one page (300–400 words)
7. GRAMMAR PROOF every sentence: check commas, apostrophes, subject-verb agreement, run-on sentences
8. Use the candidate's name, the company name, and the role title multiple times — never generic
9. Mirror language from the job description where natural
10. Tone: Confident, specific, human. Not robotic. Not sycophantic.

Return ONLY the letter body text (no subject line, no "Dear...", no meta-commentary).
The output should be clean paragraphs that can be placed directly into a letter."""


# ═══════════════════════════════════════════════════════════════
# LINKEDIN SUMMARY PROMPT
# ═══════════════════════════════════════════════════════════════

LINKEDIN_PROMPT = """You are a LinkedIn profile expert. You have helped 5,000+ professionals build profiles that attract recruiters.

Write a compelling LinkedIn About section for this candidate.

RULES:
1. First line must be a hook that makes a recruiter keep reading (not "Experienced professional with X years")
2. 3–4 short paragraphs, 220–260 words total (LinkedIn's sweet spot)
3. Include the candidate's top 3 achievements with numbers
4. End with a clear statement of what they are looking for or what value they bring
5. Naturally weave in 5–8 keywords from their industry for search visibility
6. Write in first person, conversational professional tone
7. No buzzwords: "synergy", "leverage", "passionate", "guru", "ninja", "rockstar"
8. Final line: A call to action (e.g., "Open to new opportunities in X — feel free to connect.")
9. Grammar-proof every sentence

Return only the LinkedIn About section text — clean and ready to paste."""


# ═══════════════════════════════════════════════════════════════
# GRAMMAR & TONE AUDIT PROMPT
# ═══════════════════════════════════════════════════════════════

GRAMMAR_AUDIT_PROMPT = """You are a meticulous professional editor and grammar expert.

Perform a deep grammar, style, and tone audit of the provided text.

Return EXACTLY this format:

GRAMMAR SCORE: XX/100

ERRORS FOUND:
• [Exact quote of the error] → [Correction] — [Brief explanation of the rule]
(List every grammar, spelling, punctuation, and usage error found. If none: "No grammar errors detected.")

STYLE ISSUES:
• [Style problem] → [Suggested fix]
(e.g., passive voice, weak verbs, redundant phrases, inconsistent tense)

READABILITY:
• Sentence variety: [assessment]
• Average sentence length: [short/medium/long — and whether appropriate]
• Vocabulary level: [assessment]

TONE CONSISTENCY: [assessment of whether tone is consistent throughout]

CORRECTED VERSION:
[Full corrected text with all errors fixed and style improved. Preserve the original meaning and structure. Only fix what's wrong.]

Be thorough. Professional documents must be flawless."""


# ═══════════════════════════════════════════════════════════════
# DATA & UTILITY HELPERS
# ═══════════════════════════════════════════════════════════════

def load_history():
    if os.path.exists(HISTORY_FILE):
        try:
            with open(HISTORY_FILE) as f: return json.load(f)
        except Exception: return []
    return []

def save_session(role, mode, difficulty, qa_pairs, scores):
    history = load_history()
    history.append({
        "id": len(history) + 1,
        "date": datetime.now().strftime("%Y-%m-%d %H:%M"),
        "role": role, "mode": mode, "difficulty": difficulty,
        "qa_pairs": qa_pairs, "scores": scores,
        "avg_score": round(sum(scores) / len(scores), 1) if scores else 0,
    })
    with open(HISTORY_FILE, "w") as f:
        json.dump(history, f, indent=2)

def extract_interview_score(text):
    m = re.search(r"SCORE:\s*([1-9]|10)\s*/\s*10", text)
    if m: return int(m.group(1))
    m = re.search(r"\b([1-9]|10)\s*/\s*10\b", text)
    if m: return int(m.group(1))
    return None

def extract_match_score(text):
    m = re.search(r"MATCH SCORE:\s*(\d+)", text)
    return int(m.group(1)) if m else None

def extract_section_scores(text):
    scores = {}
    patterns = [
        ("experience", r"Work Experience:\s*(\d+)/25"),
        ("skills",     r"Skills & Keywords:\s*(\d+)/25"),
        ("education",  r"Education.*?:\s*(\d+)/20"),
        ("format",     r"CV Format.*?:\s*(\d+)/15"),
        ("impact",     r"Achievements.*?:\s*(\d+)/15"),
    ]
    for key, pat in patterns:
        m = re.search(pat, text)
        if m: scores[key] = int(m.group(1))
    return scores

def get_analytics(history):
    if not history: return {}
    all_scores = [s for sess in history for s in sess.get("scores", [])]
    if not all_scores: return {}
    session_avgs = [s["avg_score"] for s in history[-12:] if s.get("avg_score")]
    return {
        "total_sessions":  len(history),
        "total_answers":   len(all_scores),
        "overall_avg":     round(sum(all_scores) / len(all_scores), 1),
        "best_score":      max(all_scores),
        "worst_score":     min(all_scores),
        "above_8":         len([s for s in all_scores if s >= 8]),
        "below_5":         len([s for s in all_scores if s <= 5]),
        "session_avgs":    session_avgs,
        "all_scores":      all_scores,
        "trend":           "up" if len(session_avgs) >= 2 and session_avgs[-1] > session_avgs[0] else "flat",
        "recent_roles":    list({s["role"] for s in history[-5:]}),
    }

def get_weak_context(history):
    if not history: return ""
    scores = [s for sess in history[-5:] for s in sess.get("scores", [])]
    if not scores: return ""
    avg = sum(scores) / len(scores)
    low = len([s for s in scores if s <= 5])
    if low > 0:
        return f"Recent average: {avg:.1f}/10. {low} answers scored 5 or below. Prioritise specificity and quantified results."
    return f"Recent average: {avg:.1f}/10. Push harder for measurable outcomes in every answer."

def extract_cv_text(uploaded_file):
    name = uploaded_file.name.lower()
    try:
        raw = uploaded_file.getvalue()
    except Exception:
        raw = uploaded_file.read()

    if name.endswith(".pdf") and PDF_SUPPORT:
        try:
            reader = PyPDF2.PdfReader(io.BytesIO(raw))
            return "\n".join(p.extract_text() or "" for p in reader.pages)
        except Exception:
            return ""

    if name.endswith(".docx") and DOCX:
        try:
            doc = DocxDocument(io.BytesIO(raw))
            return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        except Exception:
            return ""

    try:
        return raw.decode("utf-8", errors="ignore")
    except Exception:
        return ""

def slugify_filename(value, fallback="document"):
    cleaned = re.sub(r"[^A-Za-z0-9]+", "_", (value or "").strip()).strip("_")
    return cleaned[:60] or fallback

def infer_company_name(role_text):
    if not role_text:
        return "the company"
    if " at " in role_text:
        return role_text.split(" at ", 1)[-1].strip() or "the company"
    return "the company"

def split_role_and_company(role_text):
    if not role_text:
        return "", ""
    if " at " in role_text:
        role, company = role_text.split(" at ", 1)
        return role.strip(), company.strip()
    return role_text.strip(), ""

def sync_career_state(cv_text="", jd_text="", target_role="", candidate_name=""):
    if cv_text:
        st.session_state._cv_raw = cv_text
    if jd_text:
        st.session_state._jd_text = jd_text
    if target_role:
        st.session_state._target_role = target_role
    if candidate_name:
        st.session_state._candidate_name = candidate_name

def generate_word_compatible_doc(title: str, body: str) -> bytes:
    safe_title = (title or "Career Document").replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
    safe_body = (body or "").replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;").replace("\n", "<br>")
    html = f"""<html>
<head>
<meta charset="utf-8">
<title>{safe_title}</title>
</head>
<body style="font-family:Calibri,Arial,sans-serif;font-size:11pt;line-height:1.6;color:#111;padding:28px">
<h1 style="font-size:18pt;margin-bottom:18px">{safe_title}</h1>
<div>{safe_body}</div>
</body>
</html>"""
    return html.encode("utf-8")

def scrape_job_url(url):
    if not SCRAPE: return "", "Web scraping library not available."
    try:
        headers = {
            "User-Agent": "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0 Safari/537.36",
            "Accept-Language": "en-US,en;q=0.9",
        }
        resp = requests.get(url, headers=headers, timeout=15)
        if resp.status_code != 200:
            return "", f"Could not fetch URL (status {resp.status_code}). Please paste the job description manually."
        soup = BeautifulSoup(resp.text, "html.parser")
        for tag in soup(["script", "style", "nav", "header", "footer", "aside", "iframe"]): tag.decompose()
        text = soup.get_text(separator="\n")
        lines = [l.strip() for l in text.splitlines() if l.strip() and len(l.strip()) > 25]
        result = "\n".join(lines[:150])
        if len(result) < 200:
            return "", "Could not extract meaningful content from this URL. Please paste the job description manually."
        return result, ""
    except Exception as e:
        return "", f"Scraping failed: {str(e)}. Please paste the job description manually."

def export_history_csv(history):
    out = io.StringIO()
    w = csv.writer(out)
    w.writerow(["Date", "Role", "Mode", "Difficulty", "Avg Score", "Total Answers", "Best Score"])
    for s in history:
        scores = s.get("scores", [])
        w.writerow([
            s.get("date",""), s.get("role",""), s.get("mode",""),
            s.get("difficulty",""), s.get("avg_score",""),
            len(scores), max(scores) if scores else "",
        ])
    return out.getvalue()


# ═══════════════════════════════════════════════════════════════
# DOCX DOCUMENT GENERATORS
# ═══════════════════════════════════════════════════════════════

def _hr(doc, hex_color="C0A060"):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single"); bot.set(qn("w:sz"), "6")
    bot.set(qn("w:space"), "1");    bot.set(qn("w:color"), hex_color)
    pBdr.append(bot); pPr.append(pBdr)
    return p

def generate_cv_docx(cv_text_rewritten: str, candidate_name: str) -> bytes:
    if not DOCX: return cv_text_rewritten.encode()
    doc = DocxDocument()

    for sec in doc.sections:
        sec.top_margin    = Cm(1.8)
        sec.bottom_margin = Cm(1.8)
        sec.left_margin   = Cm(2.2)
        sec.right_margin  = Cm(2.2)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)

    # Parse sections
    sections_raw = {}
    current_key  = None
    current_lines = []
    section_order = []
    for line in cv_text_rewritten.splitlines():
        ls = line.strip()
        m = re.match(r"^\[([A-Z &]+)\]$", ls)
        if m:
            if current_key: sections_raw[current_key] = "\n".join(current_lines).strip()
            current_key = m.group(1); current_lines = []; section_order.append(current_key)
        elif current_key:
            current_lines.append(ls)
    if current_key: sections_raw[current_key] = "\n".join(current_lines).strip()

    name    = sections_raw.get("NAME", candidate_name or "Candidate").strip()
    contact = sections_raw.get("CONTACT", "").strip()

    # Header
    hp = doc.add_paragraph()
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hp.paragraph_format.space_after = Pt(3)
    hr = hp.add_run(name.upper())
    hr.bold = True; hr.font.size = Pt(20); hr.font.name = "Calibri"
    hr.font.color.rgb = RGBColor(0x0d, 0x1a, 0x2d)

    if contact:
        parts = [p.strip() for p in re.split(r"[|\n·•]", contact) if p.strip()]
        cp = doc.add_paragraph("  |  ".join(parts))
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_after = Pt(8)
        for r in cp.runs:
            r.font.size = Pt(9); r.font.color.rgb = RGBColor(0x55, 0x65, 0x77)

    _hr(doc, "C0A060")

    skip = {"NAME", "CONTACT"}
    for sec in section_order:
        if sec in skip: continue
        content = sections_raw.get(sec, "").strip()
        if not content: continue

        # Section title
        sp = doc.add_paragraph()
        sp.paragraph_format.space_before = Pt(12)
        sp.paragraph_format.space_after  = Pt(2)
        sr = sp.add_run(sec.upper())
        sr.bold = True; sr.font.size = Pt(9.5); sr.font.name = "Calibri"
        sr.font.color.rgb = RGBColor(0xC0, 0xA0, 0x60); sr.font.all_caps = True

        _hr(doc, "E8D8A8")

        for line in content.splitlines():
            line = line.strip()
            if not line: continue

            # Job title line
            is_role = bool(re.match(r"^[A-Z].{2,}(—|–|-{1,2}).{2,}\(\d{4}", line) or
                          re.match(r"^[A-Z].{5,}(—|–|-{1,2}).{2,}(Present|\d{4})", line))
            if is_role:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(8)
                p.paragraph_format.space_after  = Pt(2)
                r = p.add_run(line)
                r.bold = True; r.font.size = Pt(10.5)
                r.font.color.rgb = RGBColor(0x0d, 0x1a, 0x2d)

            elif line.startswith(("•", "-", "·")):
                clean = line.lstrip("•-· ").strip()
                p = doc.add_paragraph(style="List Bullet")
                p.paragraph_format.space_after = Pt(1)
                p.paragraph_format.left_indent = Inches(0.2)
                r = p.add_run(clean)
                r.font.size = Pt(10)
                r.font.color.rgb = RGBColor(0x2a, 0x2a, 0x2a)

            else:
                p = doc.add_paragraph(line)
                p.paragraph_format.space_after = Pt(3)
                for r in p.runs:
                    r.font.size = Pt(10.5)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def generate_cover_letter_docx(letter_text: str, candidate_name: str,
                                job_role: str, company: str) -> bytes:
    if not DOCX: return letter_text.encode()
    doc = DocxDocument()

    for sec in doc.sections:
        sec.top_margin    = Cm(3.0)
        sec.bottom_margin = Cm(3.0)
        sec.left_margin   = Cm(3.2)
        sec.right_margin  = Cm(3.2)

    doc.styles["Normal"].font.name  = "Garamond"
    doc.styles["Normal"].font.size  = Pt(11.5)
    doc.styles["Normal"].font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)

    # Sender block
    sender_p = doc.add_paragraph()
    sender_r = sender_p.add_run(candidate_name or "Candidate")
    sender_r.bold = True; sender_r.font.size = Pt(12)
    sender_r.font.color.rgb = RGBColor(0x0d, 0x1a, 0x2d)
    sender_p.paragraph_format.space_after = Pt(4)

    # Thin gold rule under name
    _hr(doc, "C0A060")

    # Date
    dp = doc.add_paragraph(datetime.now().strftime("%B %d, %Y"))
    dp.paragraph_format.space_before = Pt(16)
    dp.paragraph_format.space_after  = Pt(16)
    for r in dp.runs:
        r.font.size = Pt(10); r.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

    # Subject
    subj = doc.add_paragraph()
    subj.paragraph_format.space_after = Pt(20)
    sr = subj.add_run(f"Application — {job_role}")
    sr.bold = True; sr.font.size = Pt(11.5)

    # Salutation
    sal = doc.add_paragraph("Dear Hiring Manager,")
    sal.paragraph_format.space_after = Pt(14)
    for r in sal.runs: r.font.size = Pt(11.5)

    # Body
    for para in letter_text.strip().split("\n\n"):
        para = para.strip()
        if not para or para.lower().startswith("dear"): continue
        p = doc.add_paragraph(para)
        p.paragraph_format.space_after   = Pt(12)
        p.paragraph_format.alignment     = WD_ALIGN_PARAGRAPH.JUSTIFY
        for r in p.runs:
            r.font.size = Pt(11.5); r.font.name = "Garamond"

    # Closing
    close = doc.add_paragraph()
    close.paragraph_format.space_before = Pt(12)
    close.add_run("Yours sincerely,\n\n\n")
    nr = close.add_run(candidate_name or "Candidate")
    nr.bold = True; nr.font.size = Pt(11.5)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ═══════════════════════════════════════════════════════════════
# PAGE CONFIG
# ═══════════════════════════════════════════════════════════════

st.set_page_config(
    page_title="CareerCraft Pro",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ═══════════════════════════════════════════════════════════════
# WORLD-CLASS CSS
# ═══════════════════════════════════════════════════════════════

# ═══════════════════════════════════════════════════════════════
# WORLD-CLASS CSS — FIXED SIDEBAR + ANIMATIONS
# ═══════════════════════════════════════════════════════════════

st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Fraunces:ital,opsz,wght@0,9..144,300;0,9..144,400;0,9..144,500;1,9..144,300;1,9..144,400&family=DM+Mono:wght@300;400;500&family=DM+Sans:wght@300;400;500;600&display=swap');

:root {
  --ink:  #08090d; --ink2: #0f1219; --ink3: #161b26; --ink4: #1e2535;
  --bone: #f0ece4; --bone2:#e8e3da; --mist: #a8a098; --fog:  #4a5060;
  --gold: #c8a84b; --gold2:#d4b862; --gold3:#e8d090;
  --g:    #4ade80; --a:    #fbbf24; --re:   #f87171;
  --r:8px; --r2:12px; --r3:16px;
}

@keyframes fadeUp   { from{opacity:0;transform:translateY(16px)} to{opacity:1;transform:translateY(0)} }
@keyframes fadeIn   { from{opacity:0} to{opacity:1} }
@keyframes shimmer  { 0%{background-position:-200% center} 100%{background-position:200% center} }
@keyframes scaleIn  { from{opacity:0;transform:scale(0.97)} to{opacity:1;transform:scale(1)} }
@keyframes slideR   { from{opacity:0;transform:translateX(-12px)} to{opacity:1;transform:translateX(0)} }
@keyframes pulse    { 0%,100%{opacity:1} 50%{opacity:.6} }
@keyframes ringFill { from{stroke-dashoffset:340} }

*,*::before,*::after{box-sizing:border-box}

html,body,[class*="css"],.stApp{
  background-color:var(--ink)!important;
  color:var(--bone)!important;
  font-family:'DM Sans',sans-serif!important;
  font-size:15px!important;
  font-weight:300;
}

/* ── Noise overlay ── */
.stApp::after{
  content:'';position:fixed;inset:0;pointer-events:none;z-index:9999;
  background-image:url("data:image/svg+xml,%3Csvg viewBox='0 0 200 200' xmlns='http://www.w3.org/2000/svg'%3E%3Cfilter id='n'%3E%3CfeTurbulence type='fractalNoise' baseFrequency='0.85' numOctaves='4' stitchTiles='stitch'/%3E%3C/filter%3E%3Crect width='100%25' height='100%25' filter='url(%23n)' opacity='0.03'/%3E%3C/svg%3E");
  opacity:.5;
}

/* ── Sidebar ── */
section[data-testid="stSidebar"]{
  background:var(--ink2)!important;
  border-right:1px solid var(--ink4)!important;
}
section[data-testid="stSidebar"]>div{padding:2rem 1.5rem!important}
section[data-testid="stSidebar"] .stTextInput label,
section[data-testid="stSidebar"] .stSelectbox label,
section[data-testid="stSidebar"] .stFileUploader label{display:none}

/* ── Text inputs ── */
.stTextInput input, .stTextArea textarea,
[data-testid="stChatInput"] textarea {
  background:var(--ink3)!important;
  border:1px solid var(--ink4)!important;
  border-radius:var(--r)!important;
  color:var(--bone)!important;
  font-family:'DM Sans',sans-serif!important;
  font-size:0.92rem!important;
  font-weight:300!important;
  transition:border-color .2s,box-shadow .2s!important;
}
.stTextInput input:focus,.stTextArea textarea:focus{
  border-color:var(--gold)!important;
  box-shadow:0 0 0 3px rgba(200,168,75,.1)!important;
  outline:none!important;
}

/* ── Selectbox — FIXED: no border-radius on wrapper, only on inner ── */
div[data-baseweb="select"]{position:relative!important}
div[data-baseweb="select"] [data-baseweb="select-container"]{
  background:var(--ink3)!important;
  border:1px solid var(--ink4)!important;
  border-radius:var(--r)!important;
  transition:border-color .2s!important;
}
div[data-baseweb="select"] [data-baseweb="select-container"]:hover{
  border-color:var(--fog)!important;
}
div[data-baseweb="select"] [data-baseweb="select-container"] *{
  background:transparent!important;
  color:var(--bone)!important;
  font-family:'DM Sans',sans-serif!important;
  font-size:0.88rem!important;
  font-weight:300!important;
}
div[data-baseweb="select"] svg{color:var(--fog)!important}
div[data-baseweb="popover"]{
  background:var(--ink3)!important;
  border:1px solid var(--ink4)!important;
  border-radius:var(--r)!important;
  box-shadow:0 12px 40px rgba(0,0,0,.6)!important;
  animation:scaleIn .15s ease!important;
}
div[data-baseweb="menu"]{background:var(--ink3)!important}
li[role="option"]{
  color:var(--mist)!important;
  font-size:0.86rem!important;
  padding:9px 14px!important;
  transition:background .15s!important;
}
li[role="option"]:hover{background:var(--ink4)!important;color:var(--bone)!important}
li[aria-selected="true"]{color:var(--gold)!important;background:rgba(200,168,75,.06)!important}

/* ── Buttons ── */
.stButton>button{
  background:transparent!important;
  border:1px solid var(--ink4)!important;
  border-radius:var(--r)!important;
  color:var(--fog)!important;
  font-family:'DM Sans',sans-serif!important;
  font-size:0.74rem!important;
  font-weight:500!important;
  letter-spacing:.1em!important;
  text-transform:uppercase!important;
  padding:.55rem 1.1rem!important;
  transition:all .2s ease!important;
  height:auto!important;
}
.stButton>button:hover{
  border-color:var(--gold)!important;
  color:var(--gold)!important;
  background:rgba(200,168,75,.06)!important;
  transform:translateY(-1px)!important;
  box-shadow:0 4px 16px rgba(200,168,75,.12)!important;
}
.stButton>button:active{transform:translateY(0)!important}
button[kind="primary"]{
  background:var(--gold)!important;
  border:none!important;
  color:var(--ink)!important;
  font-weight:600!important;
}
button[kind="primary"]:hover{
  background:var(--gold2)!important;
  box-shadow:0 6px 24px rgba(200,168,75,.3)!important;
  transform:translateY(-1px)!important;
}

/* ── File uploader ── */
[data-testid="stFileUploader"]{
  border:1px dashed var(--ink4)!important;
  border-radius:var(--r2)!important;
  background:var(--ink2)!important;
  transition:border-color .2s!important;
}
[data-testid="stFileUploader"]:hover{border-color:var(--gold)!important}
[data-testid="stFileUploaderDropzone"] *{color:var(--fog)!important;font-size:.82rem!important}

/* ── Tabs ── */
[data-testid="stTabs"] [data-baseweb="tab-list"]{
  background:transparent!important;
  border-bottom:1px solid var(--ink4)!important;
  gap:0!important;
}
[data-testid="stTabs"] [data-baseweb="tab"]{
  background:transparent!important;
  border:none!important;
  color:var(--ink4)!important;
  font-family:'DM Sans',sans-serif!important;
  font-size:.72rem!important;
  font-weight:500!important;
  letter-spacing:.14em!important;
  text-transform:uppercase!important;
  padding:.95rem 1.6rem!important;
  border-bottom:2px solid transparent!important;
  transition:all .2s ease!important;
}
[data-testid="stTabs"] [data-baseweb="tab"]:hover{color:var(--mist)!important}
[data-testid="stTabs"] [aria-selected="true"]{
  color:var(--gold)!important;
  border-bottom:2px solid var(--gold)!important;
}

/* ── Chat ── */
[data-testid="stChatMessage"]{
  background:var(--ink2)!important;
  border:1px solid var(--ink4)!important;
  border-radius:var(--r2)!important;
  padding:1.5rem 1.75rem!important;
  margin-bottom:10px!important;
  animation:fadeUp .3s ease!important;
}
[data-testid="stChatMessage"] p,
[data-testid="stChatMessage"] li{
  font-size:.93rem!important;line-height:1.82!important;
  color:var(--bone2)!important;font-weight:300!important;
}
[data-testid="stChatInput"] textarea{
  background:var(--ink3)!important;border:1px solid var(--ink4)!important;
  border-radius:var(--r2)!important;font-size:.93rem!important;
}

/* ── Metrics ── */
[data-testid="metric-container"]{
  background:var(--ink2)!important;border:1px solid var(--ink4)!important;
  border-radius:var(--r2)!important;padding:1.3rem 1.5rem!important;
  transition:border-color .2s,transform .2s!important;
  animation:fadeUp .4s ease!important;
}
[data-testid="metric-container"]:hover{
  border-color:rgba(200,168,75,.3)!important;transform:translateY(-2px)!important;
}
[data-testid="stMetricValue"]{
  font-family:'DM Mono',monospace!important;font-size:1.7rem!important;
  font-weight:400!important;color:var(--bone)!important;letter-spacing:-.04em!important;
}
[data-testid="stMetricLabel"]{
  font-size:.68rem!important;font-weight:500!important;
  letter-spacing:.14em!important;text-transform:uppercase!important;color:var(--ink4)!important;
}

/* ── Expander ── */
[data-testid="stExpander"]{
  background:var(--ink2)!important;border:1px solid var(--ink4)!important;
  border-radius:var(--r)!important;margin-bottom:6px!important;
  transition:border-color .2s!important;
}
[data-testid="stExpander"]:hover{border-color:var(--fog)!important}
[data-testid="stExpander"] summary{
  font-size:.88rem!important;font-weight:300!important;
  color:var(--mist)!important;padding:.9rem 1.1rem!important;
}

/* ── Download button ── */
[data-testid="stDownloadButton"] button{
  background:transparent!important;border:1px solid var(--ink4)!important;
  color:var(--fog)!important;font-size:.72rem!important;
  letter-spacing:.1em!important;text-transform:uppercase!important;
  border-radius:var(--r)!important;padding:.5rem 1rem!important;
  transition:all .2s!important;
}
[data-testid="stDownloadButton"] button:hover{
  border-color:var(--gold)!important;color:var(--gold)!important;
  background:rgba(200,168,75,.06)!important;
}

/* ── Alert/Info ── */
[data-testid="stAlert"]{
  background:var(--ink2)!important;border:1px solid var(--ink4)!important;
  border-radius:var(--r)!important;font-size:.88rem!important;color:var(--fog)!important;
}

/* ── Progress bar ── */
[data-testid="stProgress"]>div>div{background:var(--gold)!important;border-radius:2px!important}

hr{border-color:var(--ink4)!important;margin:1.5rem 0!important}
::-webkit-scrollbar{width:3px;height:3px}
::-webkit-scrollbar-track{background:transparent}
::-webkit-scrollbar-thumb{background:var(--ink4);border-radius:2px}

/* ═══ Custom Components ═══ */

.wordmark{
  font-family:'Fraunces',serif;font-size:1.4rem;font-weight:400;
  color:var(--bone);letter-spacing:-.03em;line-height:1;font-style:italic;
}
.wordmark-acc{color:var(--gold);font-style:normal}
.wordmark-sub{
  font-family:'DM Sans',sans-serif;font-size:.58rem;font-weight:500;
  letter-spacing:.22em;text-transform:uppercase;color:var(--ink4);margin-top:5px;
}
.rule{height:1px;background:var(--ink4);margin:1rem 0}
.cap{
  font-size:.6rem;font-weight:500;letter-spacing:.2em;text-transform:uppercase;
  color:var(--ink4);margin:1.6rem 0 .65rem;display:block;
}

/* Session badge */
.sbadge{
  display:flex;flex-wrap:wrap;align-items:center;gap:8px;
  background:var(--ink2);border:1px solid var(--ink4);border-radius:var(--r);
  padding:10px 16px;font-size:.82rem;font-weight:300;color:var(--fog);
  margin-bottom:1.25rem;animation:fadeUp .3s ease;
}
.sbadge strong{color:var(--bone);font-weight:400}
.sep{width:3px;height:3px;border-radius:50%;background:var(--ink4);flex-shrink:0}

/* Score strip */
.sstrip{
  display:flex;align-items:stretch;gap:1px;
  background:var(--ink4);border:1px solid var(--ink4);
  border-radius:var(--r2);overflow:hidden;margin-bottom:1.5rem;
  animation:fadeUp .35s ease;
}
.scell{background:var(--ink2);flex:1;padding:1.1rem 1.4rem;transition:background .2s}
.scell:hover{background:var(--ink3)}
.slbl{font-size:.6rem;letter-spacing:.14em;text-transform:uppercase;color:var(--ink4);font-weight:500;display:block;margin-bottom:6px}
.sval{font-family:'DM Mono',monospace;font-size:1.4rem;font-weight:400;letter-spacing:-.04em}
.cg{color:#4ade80}.ca{color:#fbbf24}.cr{color:#f87171}.cgold{color:var(--gold)}.cbone{color:var(--bone)}

/* Hero */
.hero{padding:3rem 0 2rem;animation:fadeUp .5s ease}
.hero-h{
  font-family:'Fraunces',serif;font-size:3.5rem;font-weight:300;
  color:var(--bone);letter-spacing:-.05em;line-height:1.06;margin-bottom:1.25rem;
}
.hero-sub{
  font-size:1rem;font-weight:300;color:var(--fog);
  max-width:500px;line-height:1.78;margin-bottom:3rem;
}

/* Feature grid */
.fgrid{
  display:grid;grid-template-columns:repeat(3,1fr);
  gap:1px;background:var(--ink4);
  border:1px solid var(--ink4);border-radius:var(--r3);overflow:hidden;
}
.fcell{background:var(--ink);padding:1.75rem 1.5rem;transition:background .2s}
.fcell:hover{background:var(--ink2)}
.fn{font-family:'DM Mono',monospace;font-size:.65rem;color:var(--gold);letter-spacing:.1em;margin-bottom:.7rem}
.ft{font-family:'Fraunces',serif;font-size:1.05rem;color:var(--bone);margin-bottom:.45rem}
.fb{font-size:.82rem;font-weight:300;color:var(--ink4);line-height:1.65}

/* Analyser */
.atitle{font-family:'Fraunces',serif;font-size:2.2rem;font-weight:300;color:var(--bone);letter-spacing:-.04em;margin-bottom:.5rem;animation:fadeUp .4s ease}
.asub{font-size:.92rem;font-weight:300;color:var(--fog);max-width:600px;line-height:1.72;margin-bottom:2rem}

/* Ring score */
.ring-wrap{padding:1.5rem 0 1rem;display:flex;align-items:center;gap:2.5rem;animation:fadeUp .4s ease}
.ring{position:relative;width:140px;height:140px;flex-shrink:0}
.ring svg{transform:rotate(-90deg)}
.ring-text{position:absolute;inset:0;display:flex;flex-direction:column;align-items:center;justify-content:center}
.ring-num{font-family:'DM Mono',monospace;font-size:2rem;font-weight:400;letter-spacing:-.06em;line-height:1}
.ring-den{font-family:'DM Mono',monospace;font-size:.72rem;color:var(--fog);margin-top:2px}
.ring-lbl{font-size:.65rem;letter-spacing:.14em;text-transform:uppercase;margin-top:6px;font-weight:500}

/* Analysis card */
.ac{background:var(--ink2);border:1px solid var(--ink4);border-radius:var(--r2);padding:1.4rem 1.6rem;margin-bottom:10px;animation:fadeUp .4s ease}
.accap{font-size:.62rem;font-weight:500;letter-spacing:.18em;text-transform:uppercase;color:var(--gold);margin-bottom:.85rem}
.ac p,.ac li{font-size:.9rem;font-weight:300;color:var(--mist);line-height:1.75;margin:0 0 6px}
.ac strong{color:var(--bone2);font-weight:400}

/* Bar score */
.barwrap{margin:8px 0}
.barrow{display:flex;align-items:center;gap:12px;margin-bottom:7px}
.barlbl{font-size:.78rem;color:var(--fog);width:130px;flex-shrink:0}
.bartrack{flex:1;height:5px;background:var(--ink4);border-radius:3px;overflow:hidden}
.barfill{height:100%;border-radius:3px;transition:width .8s ease}
.barval{font-family:'DM Mono',monospace;font-size:.78rem;color:var(--mist);width:40px;text-align:right}

/* Keyword pills */
.kp{display:inline-block;padding:3px 11px;border-radius:20px;font-size:.74rem;font-family:'DM Mono',monospace;margin:2px 3px;transition:opacity .2s}
.kp:hover{opacity:.8}
.kpy{background:rgba(74,222,128,.08);color:#4ade80;border:1px solid rgba(74,222,128,.2)}
.kpn{background:rgba(248,113,113,.08);color:#f87171;border:1px solid rgba(248,113,113,.2)}

/* Verdict */
.verdict{
  background:var(--ink3);border:1px solid var(--ink4);
  border-left:3px solid var(--gold);border-radius:0 var(--r) var(--r) 0;
  padding:1.25rem 1.5rem;margin:1rem 0;
  font-size:.92rem;font-weight:300;color:var(--mist);line-height:1.78;
  animation:slideR .4s ease;
}

/* Editable text area (rewritten CV / cover letter) */
.edit-wrap{
  background:var(--ink2);border:1px solid var(--ink4);border-radius:var(--r2);
  padding:0;overflow:hidden;margin-bottom:1rem;
  animation:fadeUp .4s ease;
}
.edit-toolbar{
  display:flex;align-items:center;justify-content:space-between;
  padding:10px 16px;border-bottom:1px solid var(--ink4);
  background:var(--ink3);
}
.edit-toolbar-label{font-size:.65rem;font-weight:500;letter-spacing:.16em;text-transform:uppercase;color:var(--gold)}
.edit-toolbar-hint{font-size:.72rem;color:var(--ink4);font-style:italic}
.stTextArea textarea{
  font-family:'DM Mono',monospace!important;font-size:.82rem!important;
  line-height:1.7!important;color:var(--mist)!important;
  background:var(--ink2)!important;border:none!important;
  border-radius:0!important;resize:vertical!important;
  min-height:400px!important;padding:1.25rem!important;
}

/* Cover letter preview */
.letter-preview{
  background:var(--ink2);border:1px solid var(--ink4);border-radius:var(--r2);
  padding:2rem 2.5rem;line-height:1.9;
  font-size:.92rem;color:var(--bone2);font-weight:300;
  white-space:pre-wrap;animation:fadeUp .4s ease;
}

/* LinkedIn */
.li-prev{
  background:var(--ink2);border:1px solid var(--ink4);border-radius:var(--r2);
  padding:1.75rem 2rem;line-height:1.85;
  font-size:.92rem;color:var(--bone2);font-weight:300;
  animation:fadeUp .4s ease;
}

/* Guide */
.gc{background:var(--ink2);border:1px solid var(--ink4);border-radius:var(--r2);padding:1.4rem 1.6rem;margin-bottom:10px}
.gccap{font-size:.62rem;font-weight:500;letter-spacing:.18em;text-transform:uppercase;color:var(--gold);margin-bottom:.85rem}
.gc p,.gc li{font-size:.9rem;font-weight:300;color:var(--fog);line-height:1.78;margin:0 0 7px}
.gc strong{color:var(--mist);font-weight:400}
.stbl{width:100%;border-collapse:collapse}
.stbl td{font-family:'DM Mono',monospace;font-size:.82rem;padding:7px 0;border-bottom:1px solid var(--ink4);color:var(--fog)}
.stbl td:first-child{color:var(--bone);width:58px}

/* Step rows */
.steprow{display:flex;gap:12px;align-items:flex-start;margin-bottom:14px}
.stepnum{min-width:26px;height:26px;border-radius:50%;background:var(--ink4);color:var(--gold);font-family:'DM Mono',monospace;font-size:.72rem;display:flex;align-items:center;justify-content:center;flex-shrink:0;margin-top:1px}
.stt{font-size:.92rem;font-weight:400;color:var(--bone);margin-bottom:3px}
.std{font-size:.82rem;font-weight:300;color:var(--fog);line-height:1.6}

/* Info banner */
.infobanner{
  background:var(--ink2);border:1px solid var(--ink4);
  border-left:3px solid var(--gold);border-radius:0 var(--r) var(--r) 0;
  padding:1.1rem 1.4rem;margin:1rem 0;
  font-size:.88rem;font-weight:300;color:var(--fog);line-height:1.72;
  animation:slideR .3s ease;
}
.infobanner strong{color:var(--mist);font-weight:400}

/* Download row */
.dlrow{display:flex;gap:10px;flex-wrap:wrap;padding:.5rem 0 1rem}

/* Grammar score */
.gscore{font-family:'DM Mono',monospace;font-size:3rem;font-weight:400;letter-spacing:-.06em;line-height:1;margin-bottom:.4rem}

/* Focus areas */
.focuscard{background:var(--ink2);border:1px solid var(--ink4);border-radius:var(--r2);padding:1.4rem 1.6rem}
.focuscard .flbl{font-size:.62rem;letter-spacing:.16em;text-transform:uppercase;color:var(--ink4);margin-bottom:8px;font-weight:500}
.focuscard .fnum{font-family:'DM Mono',monospace;font-size:2.8rem;font-weight:400;letter-spacing:-.05em;line-height:1}
.focuscard .fsub{font-size:.8rem;color:var(--fog);margin-top:5px}

/* Perf bars */
.pbrow{display:flex;align-items:center;gap:12px;margin-bottom:7px}
.pblbl{font-size:.78rem;color:var(--fog);width:80px;text-align:right;flex-shrink:0}
.pbtrack{flex:1;height:6px;background:var(--ink4);border-radius:3px;overflow:hidden}
.pbfill{height:100%;border-radius:3px;transition:width .9s ease}
.pbval{font-family:'DM Mono',monospace;font-size:.78rem;color:var(--mist);width:40px}

/* Editable section actions */
.action-row{display:flex;gap:8px;flex-wrap:wrap;margin-top:.75rem}
</style>
""", unsafe_allow_html=True)

# ═══════════════════════════════════════════════════════════════
# SESSION STATE
# ═══════════════════════════════════════════════════════════════

_defaults = {
    "messages":[], "scores":[], "qa_pairs":[],
    "session_role":"", "session_mode":"full", "session_diff":"mid",
    "session_industry":"Technology / Software",
    "cv_text":"", "session_saved":False, "session_start":None,
    "analysis_result":None, "match_score":None, "section_scores":{},
    "rewritten_cv":None, "rewritten_cv_edited":None,
    "cover_letter":None, "cover_letter_edited":None,
    "linkedin_summary":None, "linkedin_summary_edited":None,
    "cv_grammar_result":None, "cover_letter_grammar_result":None,
    "_cv_raw":"", "_jd_text":"", "_target_role":"", "_candidate_name":"",
}
for k, v in _defaults.items():
    if k not in st.session_state:
        st.session_state[k] = v

client = Groq(api_key=GROQ_API_KEY)

# ═══════════════════════════════════════════════════════════════
# SIDEBAR
# ═══════════════════════════════════════════════════════════════

with st.sidebar:
    st.markdown('<div class="wordmark"><em>Career</em><span class="wordmark-acc">Craft</span></div><div class="wordmark-sub">Pro Edition</div>', unsafe_allow_html=True)
    st.markdown('<div class="rule"></div>', unsafe_allow_html=True)

    st.markdown('<span class="cap">Interview Session</span>', unsafe_allow_html=True)
    role_input   = st.text_input("role_inp", placeholder="Senior PM at Stripe", label_visibility="collapsed")
    industry_sel = st.selectbox("ind_sel", INDUSTRIES, label_visibility="collapsed")
    c1, c2 = st.columns(2)
    with c1: mode_sel = st.selectbox("mode_sel", list(MODES.keys()), label_visibility="collapsed")
    with c2: diff_sel = st.selectbox("diff_sel", list(DIFFICULTIES.keys()), label_visibility="collapsed")
    q_num = st.selectbox("q_sel", ["5 questions","8 questions","10 questions","15 questions"], label_visibility="collapsed")

    st.markdown('<span class="cap">Resume / CV  <span style="font-size:0.6rem;color:var(--ink4);">(optional for interview)</span></span>', unsafe_allow_html=True)
    cv_file = st.file_uploader("cv_up", type=["pdf","txt","docx"], label_visibility="collapsed")
    if cv_file:
        st.markdown(f'<div style="font-size:.76rem;color:var(--gold);margin-top:3px;">{cv_file.name} — ready</div>', unsafe_allow_html=True)

    st.markdown('<div style="height:.75rem"></div>', unsafe_allow_html=True)
    start_btn = st.button("Start Interview Session", type="primary", use_container_width=True)

    st.markdown('<span class="cap">Quick Actions</span>', unsafe_allow_html=True)
    if st.button("More Questions", use_container_width=True):
        if st.session_state.messages:
            st.session_state.messages.append({"role":"user","content":f"Give me {q_num.split()[0]} more tailored questions for this role."})
            st.rerun()
    if st.button("Coaching Tips", use_container_width=True):
        if st.session_state.messages:
            st.session_state.messages.append({"role":"user","content":"Give me 3 highly specific, actionable coaching tips to improve my answers for this exact role."})
            st.rerun()
    if st.button("Company Research", use_container_width=True):
        if st.session_state.messages and st.session_state.session_role:
            st.session_state.messages.append({"role":"user","content":f"Give me 5 specific talking points about the company culture, values, and recent news for: {st.session_state.session_role}"})
            st.rerun()

    history  = load_history()
    analytics = get_analytics(history)
    if analytics:
        st.markdown('<span class="cap">Your Progress</span>', unsafe_allow_html=True)
        ca, cb = st.columns(2)
        ca.metric("Sessions", analytics["total_sessions"])
        cb.metric("Avg Score", f"{analytics['overall_avg']}/10")
        st.download_button("Export Progress",
            data=export_history_csv(history),
            file_name=f"careercraft_{datetime.now().strftime('%Y%m%d')}.csv",
            mime="text/csv", use_container_width=True)
    if history:
        st.markdown('<div style="height:.25rem"></div>', unsafe_allow_html=True)
        if st.button("Clear History", use_container_width=True):
            if os.path.exists(HISTORY_FILE): os.remove(HISTORY_FILE)
            st.rerun()


# ═══════════════════════════════════════════════════════════════
# START SESSION
# ═══════════════════════════════════════════════════════════════

if start_btn:
    if not role_input.strip():
        st.error("Please enter a job role to begin.")
        st.stop()
    if st.session_state.scores and not st.session_state.session_saved:
        save_session(st.session_state.session_role, st.session_state.session_mode,
            st.session_state.session_diff, st.session_state.qa_pairs, st.session_state.scores)
    mode_key = MODES[mode_sel]
    diff_key = DIFFICULTIES[diff_sel]
    st.session_state.update({
        "messages":[], "scores":[], "qa_pairs":[], "session_saved":False,
        "session_role":role_input, "session_mode":mode_key, "session_diff":diff_key,
        "session_industry":industry_sel,
        "session_start":datetime.now().strftime("%H:%M"), "cv_text":"",
    })
    cv_text = extract_cv_text(cv_file) if cv_file else ""
    st.session_state.cv_text = cv_text[:4500]
    weak_ctx = get_weak_context(load_history())
    n = q_num.split()[0]
    opening = f"I am preparing for: {role_input}\nIndustry: {industry_sel}\nGenerate exactly {n} questions."
    if cv_text: opening += f"\n\nMy resume:\n{cv_text[:3500]}"
    if weak_ctx: opening += f"\n\nPrevious performance: {weak_ctx}"
    st.session_state.messages.append({"role":"user","content":opening})

# ═══════════════════════════════════════════════════════════════
# TABS
# ═══════════════════════════════════════════════════════════════

tabs = st.tabs(["  Interview  ","  CV Analyser  ","  Cover Letter  ","  LinkedIn  ","  Analytics  ","  History  ","  Guide  "])
t_interview, t_analyser, t_cover, t_linkedin, t_analytics, t_history, t_guide = tabs


# ═══════════════════════════════════════════════════════════════
# TAB 1 — INTERVIEW
# ═══════════════════════════════════════════════════════════════

with t_interview:
    if not st.session_state.messages:
        st.markdown("""
        <div class="hero">
          <div class="hero-h">Land your<br>dream role.<br><span style="color:var(--gold);font-style:italic;">Every time.</span></div>
          <div class="hero-sub">AI-graded mock interviews with STAR feedback, rewritten answers, and real coaching. Set your role in the sidebar to begin.</div>
          <div class="fgrid">
            <div class="fcell"><div class="fn">01 — Interview</div><div class="ft">Real-Grade Practice</div><div class="fb">6 modes. STAR grading with scored feedback, rewritten answers, and follow-up questions after every response.</div></div>
            <div class="fcell"><div class="fn">02 — CV Analyser</div><div class="ft">ATS Score + Rewrite</div><div class="fb">Match score out of 100. Keyword gaps. Section-by-section scoring. Full CV rewrite with grammar proof — no interview needed.</div></div>
            <div class="fcell"><div class="fn">03 — Cover Letter</div><div class="ft">3 Tailored Versions</div><div class="fb">Standard, Bold, and Personalised styles. Grammar-proofed. Editable inline. Downloads as polished .docx files.</div></div>
          </div>
          <div style="font-size:.82rem;color:var(--ink4);margin-top:1.25rem;">Interview: set role in sidebar → <span style="color:var(--gold);">Start Interview Session</span> &nbsp;·&nbsp; CV Analyser: go directly to that tab</div>
        </div>""", unsafe_allow_html=True)
    else:
        diff_lbl = next((k.split("(")[0].strip() for k,v in DIFFICULTIES.items() if v==st.session_state.session_diff),"")
        cv_ind   = " · Resume loaded" if st.session_state.cv_text else ""
        st.markdown(f"""
        <div class="sbadge">
          <strong>{st.session_state.session_role}</strong>
          <div class="sep"></div>{st.session_state.session_mode.title()}
          <div class="sep"></div>{diff_lbl}
          <div class="sep"></div>{st.session_state.session_industry}{cv_ind}
          <div class="sep"></div>{st.session_state.session_start or "—"}
        </div>""", unsafe_allow_html=True)

        if st.session_state.scores:
            sc = st.session_state.scores; avg=sum(sc)/len(sc); last=sc[-1]
            lc="cg" if last>=8 else("ca" if last>=6 else"cr")
            ac="cg" if avg>=8  else("ca" if avg>=6  else"cr")
            dh=""
            if len(sc)>=2:
                d=sc[-1]-sc[-2]; dc="cg" if d>=0 else"cr"
                dh=f'<span class="{dc}" style="font-family:DM Mono,monospace;font-size:.76rem;margin-left:8px;">{"+" if d>=0 else""}{d}</span>'
            st.markdown(f"""
            <div class="sstrip">
              <div class="scell"><span class="slbl">Last Answer</span><span class="sval {lc}">{last}/10</span>{dh}</div>
              <div class="scell"><span class="slbl">Session Avg</span><span class="sval {ac}">{avg:.1f}/10</span></div>
              <div class="scell"><span class="slbl">Graded</span><span class="sval cbone">{len(sc)}</span></div>
              <div class="scell"><span class="slbl">Best</span><span class="sval cg">{max(sc)}/10</span></div>
            </div>""", unsafe_allow_html=True)

        for i, msg in enumerate(st.session_state.messages):
            if i==0 and msg["role"]=="user":
                with st.chat_message("user",avatar=None):
                    st.markdown(f"**{st.session_state.session_role}**{cv_ind}")
            else:
                with st.chat_message(msg["role"],avatar=None):
                    st.markdown(msg["content"])

        if st.session_state.messages and st.session_state.messages[-1]["role"]=="user":
            sys_p = build_interview_prompt(st.session_state.session_mode, st.session_state.session_diff,
                        st.session_state.session_industry, st.session_state.cv_text)
            with st.chat_message("assistant",avatar=None):
                with st.spinner(""):
                    try:
                        resp = client.chat.completions.create(
                            model=GROQ_MODEL,
                            messages=[{"role":"system","content":sys_p}]+st.session_state.messages,
                            max_tokens=1800, temperature=0.7)
                        reply = resp.choices[0].message.content
                        st.markdown(reply)
                    except Exception as e:
                        reply = f"Error: {e}"; st.error(reply)
            st.session_state.messages.append({"role":"assistant","content":reply})
            score = extract_interview_score(reply)
            if score is not None:
                st.session_state.scores.append(score)
                prev = st.session_state.messages[-2]["content"] if len(st.session_state.messages)>=2 else ""
                st.session_state.qa_pairs.append({"answer":prev[:400],"score":score,"date":datetime.now().strftime("%Y-%m-%d %H:%M")})
                st.rerun()

        if user_in := st.chat_input("Type your answer, ask for more questions, or request tips..."):
            st.session_state.messages.append({"role":"user","content":user_in})
            st.rerun()

        if st.session_state.scores and not st.session_state.session_saved:
            st.markdown('<div class="rule"></div>', unsafe_allow_html=True)
            cl, cr = st.columns([3,1])
            with cl:
                avg_d = sum(st.session_state.scores)/len(st.session_state.scores)
                st.markdown(f'<div style="font-size:.82rem;color:var(--ink4);padding-top:8px;">{len(st.session_state.scores)} answers graded &nbsp;·&nbsp; avg <span style="font-family:DM Mono,monospace;color:var(--gold);">{avg_d:.1f}/10</span></div>', unsafe_allow_html=True)
            with cr:
                if st.button("Save Session", type="primary", use_container_width=True):
                    save_session(st.session_state.session_role, st.session_state.session_mode,
                        st.session_state.session_diff, st.session_state.qa_pairs, st.session_state.scores)
                    st.session_state.session_saved = True; st.rerun()


# ═══════════════════════════════════════════════════════════════
# TAB 2 — CV ANALYSER  (works standalone — NO interview needed)
# ═══════════════════════════════════════════════════════════════

with t_analyser:
    st.markdown("""
    <div style="padding:.5rem 0 1.5rem">
      <div class="atitle">CV Analyser</div>
      <div class="asub">Upload your CV and job details for a deep analysis — ATS score, keyword gaps, section scoring, grammar audit, and a complete optimised rewrite. No interview session needed.</div>
    </div>""", unsafe_allow_html=True)

    col_l, col_r = st.columns([1,1], gap="large")
    with col_l:
        st.markdown('<span class="cap">Your CV / Resume</span>', unsafe_allow_html=True)
        a_cv = st.file_uploader("a_cv_up", type=["pdf","txt","docx"], label_visibility="collapsed", key="akey_cv")
        if a_cv: st.markdown(f'<div style="font-size:.8rem;color:#4ade80;margin-top:5px;">{a_cv.name} — loaded</div>', unsafe_allow_html=True)

        st.markdown('<span class="cap">Target Role</span>', unsafe_allow_html=True)
        a_role = st.text_input("a_role_i", placeholder="Data Scientist at DeepMind", label_visibility="collapsed", key="akey_role")

        st.markdown('<span class="cap">Your Full Name  <span style="color:var(--ink4);font-size:.6rem;">(for CV & letter)</span></span>', unsafe_allow_html=True)
        a_name = st.text_input("a_name_i", placeholder="Abid Hasan", label_visibility="collapsed", key="akey_name")

        st.markdown('<span class="cap">CV Rewrite Tone</span>', unsafe_allow_html=True)
        cv_tone = st.selectbox("cv_tone_s", CV_TONES, label_visibility="collapsed", key="akey_tone")

    with col_r:
        st.markdown('<span class="cap">Job Description</span>', unsafe_allow_html=True)
        jd_mode = st.radio("jd_mode_r", ["Paste text","Enter URL"], horizontal=True, label_visibility="collapsed", key="akey_jm")
        if jd_mode == "Paste text":
            jd_paste = st.text_area("jd_paste_a", placeholder="Paste the complete job description here...", height=200, label_visibility="collapsed", key="akey_jp")
            jd_url = ""
        else:
            jd_url = st.text_input("jd_url_i", placeholder="https://jobs.company.com/role/...", label_visibility="collapsed", key="akey_ju")
            jd_paste = ""
            if jd_url:
                if not SCRAPE:
                    st.markdown('<div class="infobanner">URL scraping requires <strong>requests</strong> and <strong>beautifulsoup4</strong>. Run: <code>pip3 install requests beautifulsoup4</code></div>', unsafe_allow_html=True)
                else:
                    st.markdown('<div style="font-size:.78rem;color:var(--ink4);margin-top:5px;">Job page will be scraped when you run analysis.</div>', unsafe_allow_html=True)

    st.markdown('<div style="height:1rem"></div>', unsafe_allow_html=True)
    bc1, bc2, bc3, bc4, bc5, _ = st.columns([1.1,1.1,1.1,1.2,1.25,1.3])
    with bc1: btn_analyse  = st.button("Run Analysis",    type="primary", key="btn_run")
    with bc2: btn_rewrite  = st.button("Rewrite CV",      key="btn_rw")
    with bc3: btn_grammar  = st.button("Grammar Audit",   key="btn_gr")
    with bc4: btn_li_cv    = st.button("LinkedIn Summary",key="btn_li")
    with bc5: btn_cover_cv = st.button("Cover Letter",    key="btn_cover_cv")

    st.markdown('<div class="rule"></div>', unsafe_allow_html=True)

    # ── Handlers ─────────────────────────────────────────────

    cv_raw_input = extract_cv_text(a_cv) if a_cv else st.session_state._cv_raw
    resolved_jd = jd_paste.strip()
    scraped_jd = ""
    scrape_err = ""
    jd_action_requested = btn_analyse or btn_rewrite or btn_cover_cv
    if jd_mode == "Enter URL" and jd_url.strip() and jd_action_requested:
        with st.spinner("Scraping job page..."):
            scraped_jd, scrape_err = scrape_job_url(jd_url.strip())
        if scrape_err:
            st.warning(scrape_err)
        resolved_jd = scraped_jd

    current_role = a_role.strip() or st.session_state._target_role
    current_name = a_name.strip() or st.session_state._candidate_name
    sync_career_state(cv_raw_input, resolved_jd, current_role, current_name)

    if st.session_state._candidate_name or st.session_state._target_role or st.session_state._cv_raw:
        cv_status = "Loaded" if st.session_state._cv_raw else "Not loaded"
        jd_status = "Loaded" if st.session_state._jd_text else "Missing"
        st.markdown(
            f"""
            <div class="infobanner">
              <strong>Career workspace</strong><br>
              Name: {st.session_state._candidate_name or "Not set"} &nbsp;·&nbsp;
              Role: {st.session_state._target_role or "Not set"} &nbsp;·&nbsp;
              CV: {cv_status} &nbsp;·&nbsp;
              JD: {jd_status}
            </div>
            """,
            unsafe_allow_html=True,
        )

    if btn_analyse:
        if not cv_raw_input:
            st.error("Please upload your CV.")
        elif not current_role:
            st.error("Please enter the target role.")
        else:
            jd = resolved_jd.strip()
            if not jd:
                st.error("Please provide a job description.")
            else:
                sync_career_state(cv_raw_input, jd, current_role, current_name)
                st.session_state.rewritten_cv  = None
                st.session_state.rewritten_cv_edited = None
                st.session_state.cover_letter  = None
                st.session_state.cover_letter_edited = None
                st.session_state.cv_grammar_result = None
                st.session_state.cover_letter_grammar_result = None
                st.session_state.linkedin_summary = None
                st.session_state.linkedin_summary_edited = None
                with st.spinner("Running deep CV analysis..."):
                    try:
                        resp = client.chat.completions.create(
                            model=GROQ_MODEL,
                            messages=[
                                {"role":"system","content":CV_ANALYSIS_PROMPT},
                                {"role":"user","content":f"Target role: {current_role}\n\nJOB DESCRIPTION:\n{jd[:3000]}\n\nCANDIDATE CV:\n{cv_raw_input[:4000]}"}
                            ],
                            max_tokens=1800, temperature=0.3)
                        result = resp.choices[0].message.content
                        st.session_state.analysis_result = result
                        st.session_state.match_score     = extract_match_score(result)
                        st.session_state.section_scores  = extract_section_scores(result)
                    except Exception as e:
                        st.error(f"Analysis failed: {e}")

    if btn_rewrite:
        if not st.session_state._cv_raw:
            st.warning("Please upload your CV first.")
        else:
            jd_for_rewrite = st.session_state._jd_text or resolved_jd.strip()
            sync_career_state(st.session_state._cv_raw, jd_for_rewrite, current_role, current_name)
            with st.spinner("Rewriting and grammar-proofing your CV..."):
                try:
                    resp = client.chat.completions.create(
                        model=GROQ_MODEL,
                        messages=[
                            {"role":"system","content":build_rewrite_prompt(cv_tone)},
                            {"role":"user","content":f"Target role: {st.session_state._target_role or 'Target role not provided'}\n\nJOB DESCRIPTION:\n{(jd_for_rewrite or 'No job description provided.')[:2500]}\n\nORIGINAL CV:\n{st.session_state._cv_raw[:4000]}"}
                        ],
                        max_tokens=2500, temperature=0.4)
                    rw = resp.choices[0].message.content
                    st.session_state.rewritten_cv = rw
                    st.session_state.rewritten_cv_edited = rw
                except Exception as e:
                    st.error(f"Rewrite failed: {e}")

    if btn_grammar:
        cv_for_audit = st.session_state.rewritten_cv_edited or st.session_state._cv_raw
        if not cv_for_audit:
            st.warning("Upload your CV first.")
        else:
            with st.spinner("Running grammar and style audit..."):
                try:
                    resp = client.chat.completions.create(
                        model=GROQ_MODEL,
                        messages=[
                            {"role":"system","content":GRAMMAR_AUDIT_PROMPT},
                            {"role":"user","content":f"Audit this CV:\n\n{cv_for_audit[:4000]}"}
                        ],
                        max_tokens=1500, temperature=0.2)
                    st.session_state.cv_grammar_result = resp.choices[0].message.content
                except Exception as e:
                    st.error(f"Grammar audit failed: {e}")

    if btn_li_cv:
        li_source = st.session_state.rewritten_cv_edited or st.session_state._cv_raw
        if not li_source:
            st.warning("Upload your CV first.")
        else:
            with st.spinner("Crafting LinkedIn About section..."):
                try:
                    resp = client.chat.completions.create(
                        model=GROQ_MODEL,
                        messages=[
                            {"role":"system","content":LINKEDIN_PROMPT},
                            {"role":"user","content":f"Target: {st.session_state._target_role or current_role}\nCV:\n{li_source[:3500]}"}
                        ],
                        max_tokens=700, temperature=0.6)
                    st.session_state.linkedin_summary = resp.choices[0].message.content
                    st.session_state.linkedin_summary_edited = resp.choices[0].message.content
                except Exception as e:
                    st.error(f"LinkedIn failed: {e}")

    if btn_cover_cv:
        if not st.session_state._cv_raw:
            st.warning("Please upload your CV first.")
        elif not (st.session_state._target_role or current_role):
            st.warning("Please enter the target role before generating a cover letter.")
        else:
            jd_for_letter = st.session_state._jd_text or resolved_jd.strip()
            company = infer_company_name(st.session_state._target_role or current_role)
            with st.spinner("Writing your cover letter..."):
                try:
                    resp = client.chat.completions.create(
                        model=GROQ_MODEL,
                        messages=[
                            {"role":"system","content":build_cover_letter_prompt("Standard Professional")},
                            {"role":"user","content":f"Candidate: {st.session_state._candidate_name or current_name or 'the candidate'}\nRole: {st.session_state._target_role or current_role}\nCompany: {company}\n\nJOB DESCRIPTION:\n{(jd_for_letter or 'No job description provided.')[:2800]}\n\nCV:\n{(st.session_state.rewritten_cv_edited or st.session_state._cv_raw)[:2800]}"}
                        ],
                        max_tokens=1000, temperature=0.65)
                    letter = resp.choices[0].message.content
                    st.session_state.cover_letter = letter
                    st.session_state.cover_letter_edited = letter
                except Exception as e:
                    st.error(f"Cover letter failed: {e}")

    # ── Display analysis ──────────────────────────────────────

    if st.session_state.analysis_result:
        ms   = st.session_state.match_score or 0
        sc_c = "#4ade80" if ms>=75 else("#fbbf24" if ms>=50 else"#f87171")
        sc_l = "Strong Match" if ms>=75 else("Moderate Match" if ms>=50 else"Needs Work")
        circ = int(2*3.14159*54)
        off  = circ*(1-ms/100)

        st.markdown(f"""
        <div class="ring-wrap">
          <div class="ring">
            <svg width="140" height="140" viewBox="0 0 140 140">
              <circle class="ring-bg" cx="70" cy="70" r="54" fill="none" stroke="var(--ink4)" stroke-width="8"/>
              <circle cx="70" cy="70" r="54" fill="none" stroke="{sc_c}" stroke-width="8"
                stroke-linecap="round" stroke-dasharray="{circ}" stroke-dashoffset="{off:.0f}"
                style="transition:stroke-dashoffset 1.2s ease;"/>
            </svg>
            <div class="ring-text">
              <div class="ring-num" style="color:{sc_c};">{ms}</div>
              <div class="ring-den">/100</div>
              <div class="ring-lbl" style="color:{sc_c};">{sc_l}</div>
            </div>
          </div>
          <div>
            <div style="font-size:.88rem;font-weight:300;color:var(--fog);max-width:420px;line-height:1.75;">
              Your CV has been scored against the job description.
              Press <strong style="color:var(--mist);">Rewrite CV</strong> to close the gap and get an optimised version you can edit and download.
            </div>
          </div>
        </div>""", unsafe_allow_html=True)

        # Section scores
        ss = st.session_state.section_scores
        if ss:
            secs = [("experience","Work Experience",25),("skills","Skills & Keywords",25),("education","Education",20),("format","Format & ATS",15),("impact","Achievements",15)]
            bars = ""
            for key, lbl, mx in secs:
                val = ss.get(key,0)
                pct = int(val/mx*100)
                clr = "#4ade80" if pct>=75 else("#fbbf24" if pct>=50 else"#f87171")
                bars += f'<div class="barrow"><div class="barlbl">{lbl}</div><div class="bartrack"><div class="barfill" style="width:{pct}%;background:{clr}"></div></div><div class="barval">{val}/{mx}</div></div>'
            st.markdown(f'<div class="ac"><div class="accap">Section Scores</div>{bars}</div>', unsafe_allow_html=True)

        # Keywords
        kf = re.search(r"KEYWORDS PRESENT:\s*(.+?)(?:\n|KEYWORDS MISSING)", st.session_state.analysis_result, re.DOTALL)
        km = re.search(r"KEYWORDS MISSING:\s*(.+?)(?:\n\n|\Z|STRENGTHS)", st.session_state.analysis_result, re.DOTALL)
        if kf or km:
            pills = ""
            if kf:
                for kw in [k.strip() for k in kf.group(1).split(",") if k.strip()][:14]:
                    pills += f'<span class="kp kpy">{kw}</span>'
            if km:
                for kw in [k.strip() for k in km.group(1).split(",") if k.strip()][:14]:
                    pills += f'<span class="kp kpn">{kw}</span>'
            st.markdown(f'<div class="ac"><div class="accap">Keyword Analysis &nbsp; <span style="font-size:.7rem;color:var(--fog);font-weight:300;letter-spacing:0;">green = present &nbsp; red = missing</span></div>{pills}</div>', unsafe_allow_html=True)

        def get_sec(text, start, ends):
            pat = rf"{re.escape(start)}:?\s*\n(.*?)(?={'|'.join(re.escape(e) for e in ends)}|\Z)"
            m = re.search(pat, text, re.DOTALL)
            return m.group(1).strip() if m else ""

        strengths = get_sec(st.session_state.analysis_result,"STRENGTHS",["CRITICAL GAPS","QUICK WINS","ATS","GRAMMAR","OVERALL"])
        gaps      = get_sec(st.session_state.analysis_result,"CRITICAL GAPS",["QUICK WINS","ATS","GRAMMAR","OVERALL"])
        wins      = get_sec(st.session_state.analysis_result,"QUICK WINS",["ATS","GRAMMAR","OVERALL"])
        ats       = get_sec(st.session_state.analysis_result,"ATS COMPATIBILITY",["GRAMMAR","OVERALL"])
        grammar_s = get_sec(st.session_state.analysis_result,"GRAMMAR & LANGUAGE AUDIT",["OVERALL"])
        verdict   = get_sec(st.session_state.analysis_result,"OVERALL VERDICT",[])

        c1, c2 = st.columns(2)
        if strengths: c1.markdown(f'<div class="ac"><div class="accap">Strengths</div><p style="white-space:pre-line">{strengths}</p></div>', unsafe_allow_html=True)
        if gaps:      c2.markdown(f'<div class="ac"><div class="accap">Critical Gaps</div><p style="white-space:pre-line">{gaps}</p></div>', unsafe_allow_html=True)
        c3, c4 = st.columns(2)
        if wins: c3.markdown(f'<div class="ac"><div class="accap">Quick Wins</div><p style="white-space:pre-line">{wins}</p></div>', unsafe_allow_html=True)
        if ats:  c4.markdown(f'<div class="ac"><div class="accap">ATS Compatibility</div><p style="white-space:pre-line">{ats}</p></div>', unsafe_allow_html=True)
        if grammar_s:
            st.markdown(f'<div class="ac"><div class="accap">Grammar & Language</div><p style="white-space:pre-line">{grammar_s}</p></div>', unsafe_allow_html=True)
        if verdict:
            st.markdown(f'<div class="verdict">{verdict}</div>', unsafe_allow_html=True)

    # ── Rewritten CV with inline editor ──────────────────────

    if st.session_state.rewritten_cv is not None:
        st.markdown('<div class="rule"></div>', unsafe_allow_html=True)
        st.markdown('<div style="font-family:Fraunces,serif;font-size:1.5rem;font-weight:300;color:var(--bone);margin-bottom:.5rem;letter-spacing:-.02em;animation:fadeUp .4s ease">Rewritten CV</div>', unsafe_allow_html=True)
        st.markdown('<div style="font-size:.82rem;color:var(--fog);margin-bottom:1rem;">Edit directly in the text area below, then download your final version.</div>', unsafe_allow_html=True)

        st.markdown('<div class="edit-wrap"><div class="edit-toolbar"><span class="edit-toolbar-label">Edit your CV</span><span class="edit-toolbar-hint">All changes saved automatically</span></div></div>', unsafe_allow_html=True)

        edited_cv = st.text_area(
            "cv_editor",
            value=st.session_state.rewritten_cv_edited or st.session_state.rewritten_cv,
            height=500,
            label_visibility="collapsed",
            key="cv_edit_area",
        )
        st.session_state.rewritten_cv_edited = edited_cv

        nv = st.session_state._candidate_name or current_name or "Candidate"
        rv = st.session_state._target_role or current_role or "Role"

        dl_c1, dl_c2, dl_c3, dl_c4 = st.columns(4)
        if DOCX:
            with dl_c1:
                docx_bytes = generate_cv_docx(edited_cv, nv)
                st.download_button(
                    "Download CV (.docx)",
                    data=docx_bytes,
                    file_name=f"CV_{slugify_filename(nv)}_{slugify_filename(rv)}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key="dl_cv_docx",
                )
        with dl_c2:
            st.download_button(
                "Download CV (.txt)",
                data=edited_cv.encode(),
                file_name=f"CV_{slugify_filename(nv)}.txt",
                mime="text/plain",
                key="dl_cv_txt",
            )
        with dl_c3:
            st.download_button(
                "Download CV (.doc)",
                data=generate_word_compatible_doc(f"{nv} CV", edited_cv),
                file_name=f"CV_{slugify_filename(nv)}.doc",
                mime="application/msword",
                key="dl_cv_doc",
            )
        with dl_c4:
            if st.button("Reset to AI Version", key="reset_cv"):
                st.session_state.rewritten_cv_edited = st.session_state.rewritten_cv
                st.rerun()

    # ── Grammar audit ─────────────────────────────────────────

    if st.session_state.cv_grammar_result:
        st.markdown('<div class="rule"></div>', unsafe_allow_html=True)
        st.markdown('<div style="font-family:Fraunces,serif;font-size:1.4rem;font-weight:300;color:var(--bone);margin-bottom:1rem;letter-spacing:-.02em">Grammar & Style Audit</div>', unsafe_allow_html=True)
        gs = re.search(r"GRAMMAR SCORE:\s*(\d+)", st.session_state.cv_grammar_result)
        if gs:
            gsv = int(gs.group(1))
            gcc = "#4ade80" if gsv>=85 else("#fbbf24" if gsv>=65 else"#f87171")
            st.markdown(f'<div class="gscore" style="color:{gcc}">{gsv}<span style="font-size:1.5rem;color:var(--fog)">/100</span></div>', unsafe_allow_html=True)
        st.markdown(f'<div class="ac"><p style="white-space:pre-line">{st.session_state.cv_grammar_result}</p></div>', unsafe_allow_html=True)

    # ── LinkedIn ──────────────────────────────────────────────

    if st.session_state.linkedin_summary:
        st.markdown('<div class="rule"></div>', unsafe_allow_html=True)
        st.markdown('<div style="font-family:Fraunces,serif;font-size:1.4rem;font-weight:300;color:var(--bone);margin-bottom:1rem;letter-spacing:-.02em">LinkedIn About Section</div>', unsafe_allow_html=True)
        li_edited = st.text_area(
            "li_editor_cv",
            value=st.session_state.linkedin_summary_edited or st.session_state.linkedin_summary,
            height=280,
            label_visibility="collapsed",
            key="li_edit_area_cv",
        )
        st.session_state.linkedin_summary_edited = li_edited
        st.markdown(f'<div class="li-prev">{li_edited}</div>', unsafe_allow_html=True)
        wc = len(li_edited.split())
        st.markdown(f'<div style="font-size:.74rem;color:var(--ink4);margin-top:6px;font-family:DM Mono,monospace">{wc} words &nbsp;·&nbsp; ideal: 220–260</div>', unsafe_allow_html=True)
        li1, li2 = st.columns(2)
        with li1:
            st.download_button("Download LinkedIn About (.txt)", data=li_edited.encode(), file_name="LinkedIn_About.txt", mime="text/plain", key="dl_li_cv")
        with li2:
            st.download_button("Download LinkedIn About (.doc)", data=generate_word_compatible_doc("LinkedIn About", li_edited), file_name="LinkedIn_About.doc", mime="application/msword", key="dl_li_cv_doc")

    if not any([st.session_state.analysis_result, st.session_state.rewritten_cv, st.session_state.cv_grammar_result]):
        st.markdown("""
        <div style="text-align:center;padding:4rem 0">
          <div style="font-family:'Fraunces',serif;font-size:1.6rem;font-weight:300;color:var(--bone);margin-bottom:.6rem;letter-spacing:-.02em">Ready for analysis.</div>
          <div style="font-size:.9rem;font-weight:300;color:var(--ink4);max-width:460px;margin:0 auto;line-height:1.75">Upload your CV, enter the role and job description above, then press <span style="color:var(--gold)">Run Analysis</span>. You can also jump straight to <span style="color:var(--gold)">Rewrite CV</span>, <span style="color:var(--gold)">Cover Letter</span>, or <span style="color:var(--gold)">LinkedIn Summary</span> without starting an interview.</div>
        </div>""", unsafe_allow_html=True)


# ═══════════════════════════════════════════════════════════════
# TAB 3 — COVER LETTER  (editable + downloadable)
# ═══════════════════════════════════════════════════════════════

with t_cover:
    default_cover_role, default_cover_company = split_role_and_company(st.session_state._target_role)

    st.markdown("""
    <div style="padding:.5rem 0 1.5rem">
      <div class="atitle">Cover Letter</div>
      <div class="asub">Three distinct styles tailored to the role and company. Grammar-proofed. Edit inline and download as a polished .docx file ready to send.</div>
    </div>""", unsafe_allow_html=True)

    st.markdown("""
    <div class="infobanner">
      This section now reuses the role, company, and job description from the <strong>CV Analyser</strong> when available. You can still manually change any field below before generating the letter.
    </div>""", unsafe_allow_html=True)

    # Inputs always visible
    cl1, cl2 = st.columns([1,1], gap="large")
    with cl1:
        st.markdown('<span class="cap">Your CV / Resume  <span style="color:var(--ink4);">(optional but recommended)</span></span>', unsafe_allow_html=True)
        cl_cv = st.file_uploader("cl_cv_up", type=["pdf","txt","docx"], label_visibility="collapsed", key="cl_cv_key")
        if cl_cv:
            st.markdown(f'<div style="font-size:.8rem;color:#4ade80;margin-top:5px;">{cl_cv.name} — loaded</div>', unsafe_allow_html=True)
        st.markdown('<span class="cap">Your Name</span>', unsafe_allow_html=True)
        cl_name = st.text_input("cl_name_i", value=st.session_state._candidate_name, placeholder="Abid Hasan", label_visibility="collapsed", key="cl_name")
        st.markdown('<span class="cap">Position  <span style="color:var(--ink4);">(prefilled from CV Analyser)</span></span>', unsafe_allow_html=True)
        cl_position = st.text_input("cl_position_i", value=default_cover_role, placeholder="Product Manager", label_visibility="collapsed", key="cl_position")
        st.markdown('<span class="cap">Company  <span style="color:var(--ink4);">(prefilled from CV Analyser)</span></span>', unsafe_allow_html=True)
        cl_company = st.text_input("cl_company_i", value=default_cover_company, placeholder="Notion", label_visibility="collapsed", key="cl_company")
        st.markdown('<span class="cap">Letter Style</span>', unsafe_allow_html=True)
        letter_style = st.radio("cl_style_r", LETTER_STYLES, label_visibility="collapsed", key="cl_style")

    with cl2:
        st.markdown('<span class="cap">Job Description  <span style="color:var(--ink4);">(paste here if not already loaded)</span></span>', unsafe_allow_html=True)
        cl_jd_extra = st.text_area("cl_jd_a",
            value=st.session_state._jd_text[:500] if st.session_state._jd_text else "",
            placeholder="Paste job description if not already loaded from CV Analyser...",
            height=100, label_visibility="collapsed", key="cl_jd")
        st.markdown('<span class="cap">Extra Context  <span style="color:var(--ink4);">(optional)</span></span>', unsafe_allow_html=True)
        cl_extra = st.text_area("cl_extra_a",
            placeholder="E.g. I've been using their product for 3 years. I met the CEO at a conference...",
            height=72, label_visibility="collapsed", key="cl_extra")

    st.markdown('<div style="height:.75rem"></div>', unsafe_allow_html=True)
    clb1, clb2, clb3, _ = st.columns([1,1,1,3])
    with clb1: gen_letter    = st.button("Generate Letter",   type="primary", key="gen_cl")
    with clb2: grammar_cl    = st.button("Grammar Check",     key="gram_cl")
    with clb3: regen_letter  = st.button("Regenerate",        key="regen_cl")

    st.markdown('<div class="rule"></div>', unsafe_allow_html=True)

    cl_cv_text = extract_cv_text(cl_cv) if cl_cv else (st.session_state.rewritten_cv_edited or st.session_state._cv_raw)
    cl_position_value = cl_position.strip() or default_cover_role
    cl_company_value = cl_company.strip() or default_cover_company
    cl_role_value = " at ".join([part for part in [cl_position_value, cl_company_value] if part]).strip()
    cl_name_value = cl_name.strip() or st.session_state._candidate_name
    cl_jd_value = cl_jd_extra.strip() or st.session_state._jd_text
    sync_career_state(cl_cv_text or "", cl_jd_value or "", cl_role_value, cl_name_value)

    st.markdown(
        f"""
        <div class="infobanner">
          <strong>Current cover-letter context</strong><br>
          Position: {cl_position_value or "Not set"} &nbsp;·&nbsp;
          Company: {cl_company_value or "Not set"} &nbsp;·&nbsp;
          JD: {"Loaded" if cl_jd_value else "Missing"} &nbsp;·&nbsp;
          CV: {"Loaded" if cl_cv_text else "Missing"}
        </div>
        """,
        unsafe_allow_html=True,
    )

    # ── Generate ──────────────────────────────────────────────

    def _gen_letter(style, name, position, company, extra):
        jd_text = cl_jd_value or ""
        cv_text = cl_cv_text or ""
        role = " at ".join([part for part in [position, company] if part]).strip() or "the target role"
        extra_note = f"\n\nExtra context: {extra}" if extra.strip() else ""
        with st.spinner("Writing your cover letter..."):
            try:
                resp = client.chat.completions.create(
                    model=GROQ_MODEL,
                    messages=[
                        {"role":"system","content":build_cover_letter_prompt(style)},
                        {"role":"user","content":f"Candidate: {name or 'the candidate'}\nPosition: {position or 'Not specified'}\nCompany: {company or 'the company'}\nRole: {role}\n\nJOB DESCRIPTION:\n{jd_text[:2800]}\n\nCV:\n{cv_text[:2800]}{extra_note}"}
                    ],
                    max_tokens=1000, temperature=0.65)
                letter = resp.choices[0].message.content
                st.session_state.cover_letter = letter
                st.session_state.cover_letter_edited = letter
            except Exception as e:
                st.error(f"Cover letter failed: {e}")

    if gen_letter or regen_letter:
        _gen_letter(letter_style, cl_name_value, cl_position_value, cl_company_value, cl_extra)
        st.rerun()

    if grammar_cl and st.session_state.cover_letter_edited:
        with st.spinner("Grammar-checking your cover letter..."):
            try:
                resp = client.chat.completions.create(
                    model=GROQ_MODEL,
                    messages=[
                        {"role":"system","content":GRAMMAR_AUDIT_PROMPT},
                        {"role":"user","content":f"Audit this cover letter:\n\n{st.session_state.cover_letter_edited}"}
                    ],
                    max_tokens=1200, temperature=0.2)
                st.session_state.cover_letter_grammar_result = resp.choices[0].message.content
            except Exception as e:
                st.error(f"Grammar check failed: {e}")

    # ── Display & Edit ────────────────────────────────────────

    if st.session_state.cover_letter is not None:
        st.markdown('<div style="font-family:Fraunces,serif;font-size:1.3rem;font-weight:300;color:var(--bone);margin-bottom:.5rem;letter-spacing:-.02em;animation:fadeUp .4s ease">Your Cover Letter</div>', unsafe_allow_html=True)
        st.markdown('<div style="font-size:.82rem;color:var(--fog);margin-bottom:1rem;">Edit directly below — all changes are preserved when you download.</div>', unsafe_allow_html=True)

        st.markdown('<div class="edit-wrap"><div class="edit-toolbar"><span class="edit-toolbar-label">Edit cover letter</span><span class="edit-toolbar-hint">Format preserved in .docx download</span></div></div>', unsafe_allow_html=True)

        edited_letter = st.text_area(
            "cl_editor",
            value=st.session_state.cover_letter_edited or st.session_state.cover_letter,
            height=450,
            label_visibility="collapsed",
            key="cl_edit_area",
        )
        st.session_state.cover_letter_edited = edited_letter

        nv = cl_name_value or "Candidate"
        rv = cl_role_value or "Role"
        cov = cl_company_value or infer_company_name(rv)

        wc = len(edited_letter.split())
        wc_clr = "var(--gold)" if 300<=wc<=420 else ("#f87171" if wc>450 else "var(--fog)")
        st.markdown(f'<div style="font-size:.74rem;font-family:DM Mono,monospace;color:{wc_clr};margin-bottom:.75rem">{wc} words &nbsp;·&nbsp; ideal: 300–400 words</div>', unsafe_allow_html=True)

        dl1, dl2, dl3, dl4 = st.columns(4)
        if DOCX:
            with dl1:
                cl_bytes = generate_cover_letter_docx(edited_letter, nv, rv, cov)
                st.download_button(
                    "Download Letter (.docx)",
                    data=cl_bytes,
                    file_name=f"CoverLetter_{slugify_filename(nv)}_{slugify_filename(cov)}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key="dl_cl_docx",
                )
        with dl2:
            st.download_button(
                "Download Letter (.txt)",
                data=edited_letter.encode(),
                file_name=f"CoverLetter_{slugify_filename(nv)}.txt",
                mime="text/plain",
                key="dl_cl_txt",
            )
        with dl3:
            st.download_button(
                "Download Letter (.doc)",
                data=generate_word_compatible_doc(f"Cover Letter - {nv}", edited_letter),
                file_name=f"CoverLetter_{slugify_filename(nv)}.doc",
                mime="application/msword",
                key="dl_cl_doc",
            )
        with dl4:
            if st.button("Reset to AI Version", key="reset_cl"):
                st.session_state.cover_letter_edited = st.session_state.cover_letter
                st.rerun()

        # Grammar audit results
        if st.session_state.cover_letter_grammar_result:
            st.markdown('<div class="rule"></div>', unsafe_allow_html=True)
            st.markdown('<div style="font-family:Fraunces,serif;font-size:1.1rem;font-weight:300;color:var(--bone);margin-bottom:.75rem">Grammar Audit</div>', unsafe_allow_html=True)
            gs2 = re.search(r"GRAMMAR SCORE:\s*(\d+)", st.session_state.cover_letter_grammar_result)
            if gs2:
                gs2v = int(gs2.group(1))
                gc2c = "#4ade80" if gs2v>=85 else("#fbbf24" if gs2v>=65 else"#f87171")
                st.markdown(f'<span style="font-family:DM Mono,monospace;font-size:2rem;color:{gc2c}">{gs2v}/100</span>', unsafe_allow_html=True)

            corr = re.search(r"CORRECTED VERSION:\s*\n(.*?)$", st.session_state.cover_letter_grammar_result, re.DOTALL)
            if corr:
                corrected_text = corr.group(1).strip()
                st.markdown('<div style="font-size:.82rem;color:var(--fog);margin:.75rem 0 .4rem">A grammar-corrected version is available below — you can apply it to your letter:</div>', unsafe_allow_html=True)
                if st.button("Apply Corrected Version to Editor", key="apply_corr"):
                    st.session_state.cover_letter_edited = corrected_text
                    st.rerun()

            st.markdown(f'<div class="ac"><p style="white-space:pre-line">{st.session_state.cover_letter_grammar_result}</p></div>', unsafe_allow_html=True)

    else:
        st.markdown("""
        <div style="text-align:center;padding:4rem 0">
          <div style="font-family:'Fraunces',serif;font-size:1.6rem;font-weight:300;color:var(--bone);margin-bottom:.6rem;letter-spacing:-.02em">Ready to write.</div>
          <div style="font-size:.9rem;color:var(--ink4);max-width:440px;margin:0 auto;line-height:1.75">Fill in your name and role above, optionally add your CV and job description, choose a style, and press <span style="color:var(--gold)">Generate Letter</span>.</div>
        </div>""", unsafe_allow_html=True)


# ═══════════════════════════════════════════════════════════════
# TAB 4 — LINKEDIN
# ═══════════════════════════════════════════════════════════════

with t_linkedin:
    st.markdown("""
    <div style="padding:.5rem 0 1.5rem">
      <div class="atitle">LinkedIn Profile</div>
      <div class="asub">Generate a keyword-optimised About section that attracts recruiters. Grammar-proofed and ready to paste directly into your profile.</div>
    </div>""", unsafe_allow_html=True)

    ll1, ll2 = st.columns([1,1], gap="large")
    with ll1:
        st.markdown('<span class="cap">Your CV / Resume</span>', unsafe_allow_html=True)
        li_cv = st.file_uploader("li_cv_up", type=["pdf","txt","docx"], label_visibility="collapsed", key="li_cv_key")
        if li_cv: st.markdown(f'<div style="font-size:.8rem;color:#4ade80;margin-top:5px;">{li_cv.name} — loaded</div>', unsafe_allow_html=True)
        st.markdown('<span class="cap">Your Name</span>', unsafe_allow_html=True)
        li_name = st.text_input("li_name_i", value=st.session_state._candidate_name, placeholder="Abid Hasan", label_visibility="collapsed", key="li_name_key")
        st.markdown('<span class="cap">Target Role / Career Goal</span>', unsafe_allow_html=True)
        li_role = st.text_input("li_role_i", value=st.session_state._target_role, placeholder="Open to Software Engineering roles in AI/ML", label_visibility="collapsed", key="li_role_key")

    with ll2:
        st.markdown('<span class="cap">Key Achievements to Feature</span>', unsafe_allow_html=True)
        li_ach = st.text_area("li_ach_a",
            placeholder="- Led a team of 12 to deliver X ahead of schedule\n- Grew revenue by 40% in 6 months\n- Published in Nature journal",
            height=130, label_visibility="collapsed", key="li_ach_key")
        st.markdown('<span class="cap">Tone</span>', unsafe_allow_html=True)
        li_tone = st.radio("li_tone_r", ["Confident Professional","Warm & Approachable","Technical Expert"], horizontal=True, label_visibility="collapsed", key="li_tone_key")

    st.markdown('<div style="height:.75rem"></div>', unsafe_allow_html=True)
    gen_li = st.button("Generate LinkedIn About", type="primary", key="gen_li_btn")
    st.markdown('<div class="rule"></div>', unsafe_allow_html=True)

    li_cv_text = extract_cv_text(li_cv) if li_cv else (st.session_state.rewritten_cv_edited or st.session_state._cv_raw)
    li_name_value = li_name.strip() or st.session_state._candidate_name
    li_role_value = li_role.strip() or st.session_state._target_role
    sync_career_state(li_cv_text or "", "", li_role_value, li_name_value)

    if gen_li:
        if not li_cv_text and not li_role_value:
            st.error("Please provide your CV or target role.")
        else:
            tone_note = {"Confident Professional":"Authoritative and accomplished.","Warm & Approachable":"Friendly, genuine, relatable.","Technical Expert":"Deep technical credibility, precise language."}.get(li_tone,"")
            ach_note  = f"\n\nFeatured achievements:\n{li_ach}" if li_ach.strip() else ""
            with st.spinner("Writing your LinkedIn About section..."):
                try:
                    resp = client.chat.completions.create(
                        model=GROQ_MODEL,
                        messages=[
                            {"role":"system","content":LINKEDIN_PROMPT+f"\n\nTONE: {tone_note}"},
                            {"role":"user","content":f"Candidate: {li_name_value or 'the candidate'}\nGoal: {li_role_value or 'career growth'}\n\nCV:\n{li_cv_text[:3500]}{ach_note}"}
                        ],
                        max_tokens=700, temperature=0.6)
                    st.session_state.linkedin_summary = resp.choices[0].message.content
                    st.session_state.linkedin_summary_edited = resp.choices[0].message.content
                except Exception as e:
                    st.error(f"LinkedIn generation failed: {e}")

    if st.session_state.linkedin_summary:
        st.markdown('<div style="font-family:Fraunces,serif;font-size:1.2rem;font-weight:300;color:var(--bone);margin-bottom:1rem;letter-spacing:-.02em">Your LinkedIn About</div>', unsafe_allow_html=True)
        li_editor = st.text_area(
            "li_editor_main",
            value=st.session_state.linkedin_summary_edited or st.session_state.linkedin_summary,
            height=260,
            label_visibility="collapsed",
            key="li_editor_main_key",
        )
        st.session_state.linkedin_summary_edited = li_editor
        st.markdown(f'<div class="li-prev">{li_editor}</div>', unsafe_allow_html=True)
        wc = len(li_editor.split())
        wc_c = "var(--gold)" if 220<=wc<=260 else "var(--fog)"
        st.markdown(f'<div style="font-size:.74rem;font-family:DM Mono,monospace;color:{wc_c};margin-top:8px">{wc} words &nbsp;·&nbsp; ideal: 220–260</div>', unsafe_allow_html=True)
        lli1, lli2 = st.columns(2)
        with lli1:
            st.download_button("Download (.txt)", data=li_editor.encode(), file_name=f"LinkedIn_{slugify_filename(li_name_value or 'About')}.txt", mime="text/plain", key="dl_li_main")
        with lli2:
            st.download_button("Download (.doc)", data=generate_word_compatible_doc("LinkedIn About", li_editor), file_name=f"LinkedIn_{slugify_filename(li_name_value or 'About')}.doc", mime="application/msword", key="dl_li_main_doc")
    else:
        st.markdown("""
        <div style="text-align:center;padding:4rem 0">
          <div style="font-family:'Fraunces',serif;font-size:1.6rem;font-weight:300;color:var(--bone);margin-bottom:.6rem;letter-spacing:-.02em">Stand out on LinkedIn.</div>
          <div style="font-size:.9rem;color:var(--ink4);max-width:430px;margin:0 auto;line-height:1.75">Upload CV or reuse the rewritten CV, enter your goal, and press Generate. You can edit the final About section and download it as <span style="color:var(--gold)">.txt</span> or <span style="color:var(--gold)">.doc</span>.</div>
        </div>""", unsafe_allow_html=True)


# ═══════════════════════════════════════════════════════════════
# TAB 5 — ANALYTICS
# ═══════════════════════════════════════════════════════════════

with t_analytics:
    history  = load_history()
    analytics = get_analytics(history)
    if not analytics:
        st.markdown('<div style="font-size:.92rem;color:var(--fog);padding:2.5rem 0">Complete and save an interview session to see analytics.</div>', unsafe_allow_html=True)
    else:
        st.markdown('<div style="font-family:Fraunces,serif;font-size:1.9rem;font-weight:300;letter-spacing:-.04em;margin-bottom:1.75rem;color:var(--bone)">Performance Overview</div>', unsafe_allow_html=True)
        c1,c2,c3,c4,c5 = st.columns(5)
        c1.metric("Sessions",    analytics["total_sessions"])
        c2.metric("Answers",     analytics["total_answers"])
        c3.metric("Avg Score",   f"{analytics['overall_avg']}/10")
        c4.metric("Best Score",  f"{analytics['best_score']}/10")
        c5.metric("Scored 8+",   analytics["above_8"])

        avgs = analytics.get("session_avgs",[])
        if len(avgs)>=2:
            st.markdown('<div style="font-family:Fraunces,serif;font-size:1.2rem;font-weight:300;color:var(--bone);margin:2rem 0 1rem;letter-spacing:-.02em">Score Trend</div>', unsafe_allow_html=True)
            bars=""
            for i, avg in enumerate(avgs):
                pct=int(avg/10*100)
                clr="#4ade80" if avg>=8 else("#fbbf24" if avg>=6 else"#f87171")
                bars+=f'<div class="pbrow"><div class="pblbl">Session {i+1}</div><div class="pbtrack"><div class="pbfill" style="width:{pct}%;background:{clr}"></div></div><div class="pbval">{avg}/10</div></div>'
            trend_c = "#4ade80" if analytics["trend"]=="up" else "var(--fog)"
            trend_l = "Improving" if analytics["trend"]=="up" else "Steady"
            st.markdown(f'<div class="ac"><div class="accap">Trend: <span style="color:{trend_c}">{trend_l}</span></div>{bars}</div>', unsafe_allow_html=True)

        all_s = analytics.get("all_scores",[])
        if all_s:
            st.markdown('<div style="font-family:Fraunces,serif;font-size:1.2rem;font-weight:300;color:var(--bone);margin:2rem 0 1rem;letter-spacing:-.02em">Score Distribution</div>', unsafe_allow_html=True)
            maxc = max((all_s.count(i) for i in range(1,11)),default=1)
            dbars=""
            for i in range(10,0,-1):
                cnt=all_s.count(i); pct=int(cnt/maxc*100) if maxc>0 else 0
                clr="#4ade80" if i>=8 else("#fbbf24" if i>=6 else"#f87171")
                dbars+=f'<div class="pbrow"><div class="pblbl">{i}/10</div><div class="pbtrack"><div class="pbfill" style="width:{pct}%;background:{clr}"></div></div><div class="pbval">{cnt}</div></div>'
            st.markdown(f'<div class="ac">{dbars}</div>', unsafe_allow_html=True)

        total=analytics["total_answers"]
        if total:
            st.markdown('<div style="font-family:Fraunces,serif;font-size:1.2rem;font-weight:300;color:var(--bone);margin:2rem 0 1rem;letter-spacing:-.02em">Focus Areas</div>', unsafe_allow_html=True)
            lp=round(analytics["below_5"]/total*100); hp=round(analytics["above_8"]/total*100)
            col1,col2=st.columns(2)
            col1.markdown(f'<div class="focuscard"><div class="flbl">Strong Answers</div><div class="fnum" style="color:#4ade80">{hp}<span style="font-size:1.2rem;color:var(--fog)">%</span></div><div class="fsub">scored 8 or above</div></div>', unsafe_allow_html=True)
            c2c="#f87171" if lp>20 else"#fbbf24"
            col2.markdown(f'<div class="focuscard"><div class="flbl">Need Improvement</div><div class="fnum" style="color:{c2c}">{lp}<span style="font-size:1.2rem;color:var(--fog)">%</span></div><div class="fsub">scored 5 or below</div></div>', unsafe_allow_html=True)

        if analytics.get("recent_roles"):
            st.markdown(f'<div style="font-size:.82rem;color:var(--ink4);margin-top:1.5rem">Recent roles: <span style="color:var(--fog)">{" · ".join(analytics["recent_roles"])}</span></div>', unsafe_allow_html=True)


# ═══════════════════════════════════════════════════════════════
# TAB 6 — HISTORY
# ═══════════════════════════════════════════════════════════════

with t_history:
    history=load_history()
    if not history:
        st.markdown('<div style="font-size:.92rem;color:var(--fog);padding:2.5rem 0">No sessions saved yet. Complete an interview and press Save Session.</div>', unsafe_allow_html=True)
    else:
        hc, hdl = st.columns([3,1])
        with hc: st.markdown(f'<div style="font-family:Fraunces,serif;font-size:1.9rem;font-weight:300;letter-spacing:-.04em;color:var(--bone)">{len(history)} Sessions</div>', unsafe_allow_html=True)
        with hdl: st.download_button("Export CSV", data=export_history_csv(history), file_name=f"careercraft_{datetime.now().strftime('%Y%m%d')}.csv", mime="text/csv", use_container_width=True)
        st.markdown('<div style="height:.75rem"></div>', unsafe_allow_html=True)
        for sess in reversed(history):
            avg=sess.get("avg_score",0); sc="#4ade80" if avg>=8 else("#fbbf24" if avg>=6 else"#f87171")
            scores_str="  ·  ".join([str(s) for s in sess.get("scores",[])])
            with st.expander(f"{sess.get('date','')}  ·  {sess.get('role','')}  ·  avg {avg}/10"):
                hc1,hc2,hc3=st.columns(3)
                hc1.markdown(f'<div style="font-size:.78rem;color:var(--ink4)">Mode<br><span style="color:var(--mist)">{sess.get("mode","—")}</span></div>', unsafe_allow_html=True)
                hc2.markdown(f'<div style="font-size:.78rem;color:var(--ink4)">Level<br><span style="color:var(--mist)">{sess.get("difficulty","—")}</span></div>', unsafe_allow_html=True)
                hc3.markdown(f'<div style="font-size:.78rem;color:var(--ink4)">Avg<br><span style="font-family:DM Mono,monospace;color:{sc}">{avg}/10</span></div>', unsafe_allow_html=True)
                if scores_str: st.markdown(f'<div style="font-size:.78rem;color:var(--ink4);margin-top:.75rem">Scores<br><span style="font-family:DM Mono,monospace;color:var(--fog)">{scores_str}</span></div>', unsafe_allow_html=True)
                for pair in sess.get("qa_pairs",[]):
                    s=pair.get("score","—"); sc2="#4ade80" if isinstance(s,int) and s>=8 else("#fbbf24" if isinstance(s,int) and s>=6 else"#f87171")
                    st.markdown(f'<div style="margin-top:7px;padding-left:12px;border-left:1px solid var(--ink4)"><span style="font-family:DM Mono,monospace;color:{sc2};font-size:.88rem">{s}/10</span><span style="margin-left:12px;font-size:.8rem;color:var(--ink4)">{pair.get("answer","")[:110]}...</span></div>', unsafe_allow_html=True)


# ═══════════════════════════════════════════════════════════════
# TAB 7 — GUIDE
# ═══════════════════════════════════════════════════════════════

with t_guide:
    col1, col2 = st.columns(2)
    with col1:
        st.markdown("""
        <div class="gc"><div class="gccap">How to Use</div>
        <div class="steprow"><div class="stepnum">1</div><div><div class="stt">CV Analyser (no interview needed)</div><div class="std">Upload CV → paste JD or URL → Run Analysis → Rewrite CV → edit inline → download .txt, .doc, or .docx</div></div></div>
        <div class="steprow"><div class="stepnum">2</div><div><div class="stt">Cover Letter</div><div class="std">Cover Letter tab → optionally upload CV → choose style → Generate → Grammar Check → edit inline → download .txt, .doc, or .docx</div></div></div>
        <div class="steprow"><div class="stepnum">3</div><div><div class="stt">Interview Practice</div><div class="std">Set role in sidebar → Start Session → answer questions → get STAR grading + rewritten answers</div></div></div>
        <div class="steprow"><div class="stepnum">4</div><div><div class="stt">LinkedIn Profile</div><div class="std">LinkedIn tab → upload CV or reuse rewritten CV → enter goal → Generate → edit inline → download .txt or .doc</div></div></div>
        </div>
        <div class="gc"><div class="gccap">The STAR Method</div>
        <p><strong>Situation</strong> — Set the scene. 1–2 sentences of context.<br>
        <strong>Task</strong> — Your specific responsibility. What was expected?<br>
        <strong>Action</strong> — What YOU did. Use "I" not "we". Most detailed part.<br>
        <strong>Result</strong> — Quantified outcome. Always use numbers.</p>
        </div>
        <div class="gc"><div class="gccap">Interview Modes</div>
        <p><strong>Full Interview</strong> — Balanced mix. Best for general practice.<br>
        <strong>Behavioural</strong> — Past experience. Most common in real interviews.<br>
        <strong>Technical</strong> — Hard skills. Use for engineering and data roles.<br>
        <strong>Situational</strong> — Hypotheticals. Common for management roles.<br>
        <strong>Speed Round</strong> — Builds conciseness under pressure.<br>
        <strong>Stress Test</strong> — Toughest mode. Prepares you for brutal interviewers.</p>
        </div>""", unsafe_allow_html=True)
    with col2:
        st.markdown("""
        <div class="gc"><div class="gccap">Interview Score Guide</div>
        <table class="stbl">
        <tr><td>9–10</td><td>Exceptional — ready for top-tier interviews</td></tr>
        <tr><td>7–8</td><td>Strong — add more specific numbers</td></tr>
        <tr><td>5–6</td><td>Decent — STAR elements missing</td></tr>
        <tr><td>3–4</td><td>Weak — vague, lacks structure</td></tr>
        <tr><td>1–2</td><td>Needs full rewrite</td></tr>
        </table>
        </div>
        <div class="gc"><div class="gccap">CV Match Score Guide</div>
        <table class="stbl">
        <tr><td>80–100</td><td>Excellent — apply with confidence</td></tr>
        <tr><td>60–79</td><td>Good — fix 2–3 keyword gaps</td></tr>
        <tr><td>40–59</td><td>Moderate — significant rewrites needed</td></tr>
        <tr><td>0–39</td><td>Weak match — major overhaul required</td></tr>
        </table>
        </div>
        <div class="gc"><div class="gccap">3-Week Strategy</div>
        <p><strong>Week 1 — Prepare</strong><br>Run CV Analyser. Fix all Quick Wins. Rewrite CV. Practice behavioural mode daily.<br><br>
        <strong>Week 2 — Strengthen</strong><br>Full Interview mode. Target 7+ average. Generate cover letter. Update LinkedIn.<br><br>
        <strong>Week 3 — Execute</strong><br>Stress Test mode. Push for 8+ average. Send rewritten CV and grammar-checked cover letter.</p>
        </div>""", unsafe_allow_html=True)
